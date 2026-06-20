# -*- coding: utf-8 -*-
"""
Generate the graduation-project thesis for the *Together* sign-language
translation system as a professional .docx document.

Run:  python scripts/build_thesis.py
Out:  Together_Thesis.docx   (in the repository root)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "Together_Thesis.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # deep blue
ACCENT2 = RGBColor(0x2E, 0x74, 0xB5)
GREY = RGBColor(0x59, 0x59, 0x59)

doc = Document()

# ───────────────────────── base styles ─────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for i, sz in [(1, 18), (2, 15), (3, 13)]:
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(sz)
    st.font.color.rgb = ACCENT
    st.font.bold = True

fig_count = [0]
tab_count = [0]
figures = []   # (label, caption)
tables = []


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = instr
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Update this field (right-click → Update Field)"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instrText); run._r.append(fldSep); run._r.append(t); run._r.append(fldEnd)


def para(text, align="justify", size=None, bold=False, italic=False, color=None,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    a = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.alignment = a
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def body(text):
    return para(text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def h1(text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_after = Pt(12)
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def _box_borders(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "8"); e.set(qn("w:color"), "B0B7C3")
        pbdr.append(e)
    pPr.append(pbdr)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "FAFBFC")
    pPr.append(shd)


def figure_placeholder(caption, height_note="Insert screenshot / diagram here"):
    fig_count[0] += 1
    label = f"Figure {fig_count[0]}"
    figures.append((label, caption))
    # Framed image-area placeholder ~2.4in tall (student inserts the figure here).
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    box.paragraph_format.space_before = Pt(6)
    box.paragraph_format.space_after = Pt(2)
    r = box.add_run("\n\n\n\n\n" + f"[ {label} — {height_note} ]" + "\n\n\n\n\n\n")
    r.italic = True; r.font.color.rgb = GREY; r.font.size = Pt(11)
    _box_borders(box)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"{label}: {caption}")
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = ACCENT
    cap.paragraph_format.space_after = Pt(10)
    return label


def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    for i, ln in enumerate(lines):
        run = p.add_run(("" if i == 0 else "\n") + ln)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # light shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    return p


def table(headers, rows, caption=None, col_widths=None):
    tab_count[0] += 1
    if caption:
        tables.append((f"Table {tab_count[0]}", caption))
        cp = doc.add_paragraph()
        cr = cp.add_run(f"Table {tab_count[0]}: {caption}")
        cr.italic = True; cr.font.size = Pt(10); cr.font.color.rgb = ACCENT
        cp.paragraph_format.space_after = Pt(3)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1F4E79")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ═══════════════════════════ TITLE PAGE ═══════════════════════════
def center_run(text, size, bold=False, color=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


center_run("German University in Cairo", 16, bold=True, color=ACCENT, space_after=2)
center_run("Faculty of Media Engineering and Technology", 13, color=GREY, space_after=2)
center_run("Computer Science and Engineering Department", 13, color=GREY, space_after=30)

center_run("Together", 34, bold=True, color=ACCENT, space_after=2)
center_run("A Real-Time, Bi-Directional Sign-Language Translation System", 17, bold=True, space_after=4)
center_run("for American (ASL) and Arabic / Egyptian (ArSL) Sign Languages", 14, italic=True, color=GREY, space_after=26)

center_run("A Graduation Project Thesis submitted to the", 12, space_after=0)
center_run("Computer Science and Engineering Department in Partial Fulfilment", 12, space_after=0)
center_run("of the Requirements for the Degree of Bachelor of Science", 12, space_after=24)

center_run("Submitted by", 12, italic=True, color=GREY, space_after=4)
for name in ["Student Name 1", "Student Name 2", "Student Name 3", "Student Name 4"]:
    center_run(name, 13, bold=True, space_after=2)

center_run("", 12, space_after=14)
center_run("Under the Supervision of", 12, italic=True, color=GREY, space_after=4)
center_run("Prof. Dr. [Supervisor Name]", 13, bold=True, space_after=2)
center_run("[Co-Supervisor / Teaching Assistant Name]", 12, space_after=24)

center_run("Cairo, Egypt — 2026", 13, bold=True, color=ACCENT, space_after=0)

# ═══════════════════════════ DECLARATION ═══════════════════════════
doc.add_page_break()
h_decl = doc.add_heading("Declaration of Originality", level=1)
body("This is to certify that:")
numbered("the work presented in this thesis comprises only our own original work "
         "towards the Bachelor of Science degree, except where due acknowledgement "
         "has been made in the text to all other material used;")
numbered("this thesis has not been submitted, in whole or in part, for any other "
         "degree or professional qualification at this or any other institution;")
numbered("all external sources, datasets, libraries, and pre-trained models used "
         "in this project have been properly cited and credited in the References.")
body("")
body("We further declare that the design, implementation, experimentation, and "
     "documentation of the Together system described in this thesis were carried "
     "out by the project team under the supervision named on the title page.")
body("")
for name in ["Student Name 1", "Student Name 2", "Student Name 3", "Student Name 4"]:
    p = doc.add_paragraph()
    p.add_run("\nSignature: ___________________________    ")
    p.add_run(name).bold = True
para("Cairo, 2026", space_before=18)

# ═══════════════════════════ ACKNOWLEDGEMENTS ═══════════════════════════
doc.add_page_break()
doc.add_heading("Acknowledgements", level=1)
body("First and foremost, we are deeply grateful to our supervisor, "
     "Prof. Dr. [Supervisor Name], whose guidance, patience, and constructive "
     "criticism shaped this project from a loose idea into a working system. The "
     "weekly discussions, the candid feedback on early prototypes, and the "
     "encouragement to pursue a genuinely bilingual, accessibility-first design "
     "were invaluable throughout the year.")
body("We extend our sincere thanks to the Faculty of Media Engineering and "
     "Technology at the German University in Cairo for providing the academic "
     "environment, computing resources, and the structured framework within which "
     "this graduation project was developed. We are equally indebted to the "
     "teaching assistants who reviewed our intermediate deliverables and helped us "
     "tighten both the engineering and the writing.")
body("We owe a special debt of gratitude to the Deaf and Hard-of-Hearing community "
     "and to the sign-language interpreters who informed our understanding of sign "
     "grammar, Topic–Comment structure, and the lived experience of communication "
     "barriers. Their perspective is the reason this project exists, and it is to "
     "them that we hope the work is ultimately useful.")
body("Finally, we thank our families and friends for their unwavering support, for "
     "tolerating late nights and long debugging sessions, and for believing in the "
     "value of building technology that brings people together.")

# ═══════════════════════════ ABSTRACT ═══════════════════════════
doc.add_page_break()
doc.add_heading("Abstract", level=1)
body("Communication between Deaf or Hard-of-Hearing (DHH) signers and the hearing, "
     "non-signing majority remains one of the most persistent accessibility gaps in "
     "everyday life. Existing assistive tools are typically uni-directional, "
     "monolingual, tethered to expensive hardware, or limited to isolated word "
     "lookups, and very few address Arabic sign language at all. This thesis "
     "presents Together, a real-time, bi-directional sign-language translation "
     "system that runs entirely in a standard web browser with nothing more than a "
     "webcam and a microphone. Together supports two first-class languages — "
     "American Sign Language (ASL) paired with English, and Arabic / Egyptian Sign "
     "Language (ArSL) paired with Arabic — and translates in both directions: from "
     "sign to text and speech, and from text and speech back into a signing avatar.")
body("The system is built around four coordinated pipelines. In the sign-to-text "
     "direction, MediaPipe Holistic extracts pose, hand, and face landmarks in the "
     "browser; these landmark sequences are streamed to a server that runs a "
     "250-class TensorFlow-Lite model for ASL and a custom CNN–GRU PyTorch model for "
     "the 20-class Arabic vocabulary. Recognised glosses are debounced with a "
     "vote-and-cooldown buffer and then assembled into natural sentences by a Large "
     "Language Model (LLM) that respects the Topic–Comment grammar of sign "
     "languages. In the text-to-sign direction, a sentence is converted into ordered "
     "gloss tokens annotated with non-manual markers, each token is matched to a "
     "stored landmark sequence through Sentence-BERT semantic search over a "
     "pgvector index, and out-of-vocabulary words are synthesised on the fly by "
     "fingerspelling and a Savitzky–Golay-smoothed stitching algorithm. A live "
     "two-party meeting mode connects a signer and a speaker over WebRTC, with "
     "Socket.IO signalling and per-role captions or avatar rendering.")
body("Every external capability — language modelling, text-to-speech, and "
     "speech-to-text — is routed through a pluggable provider abstraction that "
     "prefers a cloud provider (Google Gemini) but degrades gracefully to fully "
     "offline fallbacks (Ollama, pyttsx3, and Whisper) whenever a key is absent or a "
     "network call fails, so the core experience never breaks. The platform is "
     "engineered for production: JSON Web Token authentication with Argon2id "
     "password hashing, rotating refresh tokens, sliding-window rate limiting, a "
     "PostgreSQL database with the pgvector extension for indexed semantic search, "
     "and a comprehensive latency-optimisation phase that moves all blocking "
     "inference off the event loop, quantises the Arabic model to INT8, and caches "
     "repeated gloss translations.")
body("On a held-out evaluation of 250 ASL samples the recognition model achieves a "
     "Top-1 accuracy of 62.4%, and the Arabic INT8 model runs roughly 1.8× faster "
     "than its FP32 baseline on CPU while preserving accuracy. The user interface is "
     "fully bilingual, supporting both left-to-right English and right-to-left "
     "Arabic through CSS logical properties, and ships a light/dark themed design "
     "system. The result is an accessible, low-cost, install-free bridge for "
     "two-way communication between signers and speakers. This thesis documents the "
     "motivation, the linguistic and machine-learning foundations, the system "
     "architecture, the implementation of each pipeline, the security and data "
     "layer, the performance engineering, and a critical evaluation of the system, "
     "and closes with a discussion of limitations and directions for future work.")
body("")
p = doc.add_paragraph()
p.add_run("Keywords: ").bold = True
p.add_run("Sign Language Recognition; American Sign Language; Arabic Sign "
          "Language; MediaPipe Holistic; Deep Learning; CNN-GRU; TensorFlow Lite; "
          "Gloss Translation; Topic–Comment Grammar; Sentence-BERT; pgvector; "
          "WebRTC; Real-Time Systems; Accessibility; Human–Computer Interaction.")

# ═══════════════════════════ TOC ═══════════════════════════
doc.add_page_break()
doc.add_heading("Table of Contents", level=1)
add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u')

doc.add_page_break()
doc.add_heading("List of Figures", level=1)
add_field(doc.add_paragraph(), 'TOC \\h \\z \\c "Figure"')
body("(If figure entries do not appear, insert images using captions labelled "
     "“Figure N” and update this field.)")

doc.add_page_break()
doc.add_heading("List of Tables", level=1)
add_field(doc.add_paragraph(), 'TOC \\h \\z \\c "Table"')

# ═══════════════════════════ ABBREVIATIONS ═══════════════════════════
doc.add_page_break()
doc.add_heading("List of Abbreviations", level=1)
abbr = [
    ("ASL", "American Sign Language"),
    ("ArSL", "Arabic Sign Language"),
    ("DHH", "Deaf and Hard-of-Hearing"),
    ("SLR", "Sign Language Recognition"),
    ("SLT", "Sign Language Translation"),
    ("NMM", "Non-Manual Markers"),
    ("CNN", "Convolutional Neural Network"),
    ("RNN", "Recurrent Neural Network"),
    ("GRU", "Gated Recurrent Unit"),
    ("LSTM", "Long Short-Term Memory"),
    ("BiLSTM", "Bidirectional Long Short-Term Memory"),
    ("LLM", "Large Language Model"),
    ("TTS", "Text-to-Speech"),
    ("STT", "Speech-to-Text"),
    ("NLP", "Natural Language Processing"),
    ("SBERT", "Sentence-BERT"),
    ("TFLite", "TensorFlow Lite"),
    ("INT8", "8-bit Integer Quantization"),
    ("API", "Application Programming Interface"),
    ("REST", "Representational State Transfer"),
    ("JWT", "JSON Web Token"),
    ("RBAC", "Role-Based Access Control"),
    ("WebRTC", "Web Real-Time Communication"),
    ("ICE", "Interactive Connectivity Establishment"),
    ("STUN/TURN", "Session Traversal Utilities / Traversal Using Relays around NAT"),
    ("ORM", "Object-Relational Mapping"),
    ("HNSW", "Hierarchical Navigable Small World (vector index)"),
    ("CDR", "Cup-to-Disc Ratio (used only in cited related work)"),
    ("RTL / LTR", "Right-to-Left / Left-to-Right text direction"),
    ("CORS", "Cross-Origin Resource Sharing"),
    ("p50 / p95", "50th / 95th percentile latency"),
]
table(["Abbreviation", "Full Form"], abbr)

# ═══════════════════════════ CHAPTER 1 ═══════════════════════════
h1("Chapter 1: Introduction")

h2("1.1 Background and Motivation")
body("Sign languages are complete, natural human languages with their own "
     "phonology, morphology, and syntax. They are not manual transliterations of "
     "spoken languages; American Sign Language (ASL) is no more a gestural form of "
     "English than spoken Japanese is a spoken form of English. For millions of "
     "Deaf and Hard-of-Hearing (DHH) people worldwide, a sign language is the "
     "primary, fluent, and native means of communication. Yet the overwhelming "
     "majority of the hearing population does not understand any sign language, "
     "which produces a pervasive and exhausting communication barrier in precisely "
     "the situations where clear communication matters most: a hospital reception, "
     "a bank counter, a classroom, a job interview, or a conversation with a "
     "neighbour.")
body("According to the World Health Organization, over 1.5 billion people live with "
     "some degree of hearing loss, and that figure is projected to rise to roughly "
     "2.5 billion by 2050, with around 700 million people expected to require "
     "rehabilitation services. A significant subset of this population relies on "
     "sign language as a first language. In the Arab world specifically, access to "
     "sign-language interpreters is scarce, and digital tools that understand "
     "Arabic sign language are almost non-existent compared to the relatively "
     "richer ecosystem available for ASL. This asymmetry motivated us to treat "
     "Arabic / Egyptian Sign Language as a first-class citizen of the system rather "
     "than an afterthought.")
figure_placeholder("Global hearing-loss statistics and the projected growth of the "
                   "Deaf and Hard-of-Hearing population through 2050.")
body("Human interpreters remain the gold standard for sign-language communication, "
     "but they are expensive, must be booked in advance, are unavailable on demand, "
     "and raise legitimate privacy concerns when the conversation is personal — for "
     "instance, a medical or legal matter. Technology cannot and should not replace "
     "professional interpreters for high-stakes settings, but it can dramatically "
     "lower the barrier for the countless everyday micro-interactions that make up "
     "ordinary life and that are today simply lost or reduced to writing notes back "
     "and forth on a phone. It is exactly this gap that Together targets.")

h2("1.2 Sign Language Translation Defined")
body("Sign Language Translation (SLT) is the broad task of converting between a "
     "sign language and another language modality. It is useful to decompose it "
     "into the two directions the system must support and the sub-problems each "
     "entails.")
h3("1.2.1 Sign-to-Text and Sign-to-Speech")
body("In this direction the system observes a signer through a camera and produces "
     "written and spoken output in the corresponding spoken language. This is not a "
     "single classification problem. It begins with Sign Language Recognition "
     "(SLR), the perceptual task of identifying which signs are being produced from "
     "a stream of body, hand, and face motion. SLR alone yields a sequence of "
     "glosses — uppercase labels that stand for individual signs, such as "
     "STORE I GO — which are not yet grammatical sentences in the spoken language. "
     "A second, linguistic stage must therefore translate this gloss sequence, "
     "which follows sign-language word order, into a fluent sentence such as "
     "“I am going to the store.” Finally, for sign-to-speech, the sentence "
     "is rendered as audio through a text-to-speech engine.")
h3("1.2.2 Text-to-Sign and Speech-to-Sign")
body("In the reverse direction the system takes typed text or spoken audio and "
     "produces sign-language output through an animated avatar. Spoken input is "
     "first transcribed to text by a speech-to-text engine. The text sentence is "
     "then converted into ordered gloss tokens that follow sign-language grammar, "
     "each token is mapped to a stored motion sequence for the corresponding sign, "
     "and the sequences are played back, smoothly transitioning from one sign to the "
     "next. Words that have no dedicated sign are fingerspelled letter by letter.")
body("A complete, useful system must handle all four of these flows, in both "
     "languages, with low enough latency to feel conversational. Together is, to our "
     "knowledge, unusual in delivering all four flows for two languages from a "
     "single install-free web application.")

h2("1.3 Global and Local Statistics")
table(
    ["Indicator", "Figure", "Implication for the project"],
    [
        ["People with hearing loss (WHO, present)", "~1.5 billion", "Vast latent user base for assistive communication tools"],
        ["Projected by 2050", "~2.5 billion", "The problem is growing, not shrinking"],
        ["Estimated requiring rehabilitation by 2050", "~700 million", "Demand for low-cost, scalable solutions"],
        ["Certified interpreters relative to DHH population", "Severe shortage", "Human interpreting cannot scale to daily interactions"],
        ["Digital tools targeting Arabic sign language", "Very few", "A clear, underserved niche"],
    ],
    caption="Selected statistics motivating an accessible, scalable, bilingual "
            "sign-language translation system.")
body("These figures establish two facts that anchor the project. First, the "
     "population that could benefit is enormous and growing. Second, the supply of "
     "human interpreters is structurally incapable of meeting demand for ordinary, "
     "everyday communication, and the situation is far worse for Arabic sign "
     "language than for ASL. A free, browser-based tool that requires no special "
     "hardware is therefore not a luxury but a meaningful contribution to "
     "accessibility.")

h2("1.4 Root Causes of Communication Barriers")
body("It is worth being precise about why the barrier persists, because the design "
     "of Together follows directly from these causes.")
bullet("Asymmetry of language knowledge: signers learn to read and write the "
       "majority language to varying degrees, but hearing people almost never learn "
       "to sign, so the burden of bridging falls entirely on the DHH person.")
bullet("Scarcity and cost of interpreters: professional interpreting is expensive, "
       "requires advance booking, and is unavailable for spontaneous interactions.")
bullet("Hardware and cost barriers in prior technical solutions: many research "
       "systems depend on data gloves, depth cameras, or wearable sensors that are "
       "costly, intrusive, and impractical for everyday use.")
bullet("Linguistic complexity: sign languages use spatial grammar, simultaneous "
       "morphology, and non-manual markers (facial expressions, head movement) that "
       "naïve word-for-word systems cannot represent.")
bullet("Language coverage: the few consumer tools that exist overwhelmingly target "
       "ASL/English and ignore Arabic sign language entirely.")
body("Together addresses each of these directly: it uses only a commodity webcam "
     "(no special hardware), it is free and instantly available in the browser, it "
     "models sign grammar through gloss translation and a non-manual-marker layer, "
     "and it treats Arabic sign language as a first-class language with its own "
     "model, vocabulary, and right-to-left interface.")

h2("1.4b The Economic and Social Dimension")
body("The communication barrier is not only a matter of individual inconvenience; it "
     "carries a measurable economic and social cost. When Deaf and Hard-of-Hearing "
     "people cannot communicate freely, they are systematically disadvantaged in "
     "education, where instruction is overwhelmingly delivered in spoken language; in "
     "employment, where interviews and day-to-day collaboration assume hearing; and "
     "in healthcare, where the inability to communicate symptoms accurately can be "
     "dangerous. Each of these exclusions has a cost — to the individual in lost "
     "opportunity and to society in unrealised human capital. Assistive technology "
     "that lowers the barrier therefore has value far beyond the convenience of any "
     "single conversation; it is an investment in inclusion.")
body("There is also a dignity dimension that pure statistics miss. Relying on a "
     "third party to interpret every personal interaction — a doctor's visit, a "
     "disagreement, a private conversation — erodes autonomy and privacy. A tool "
     "that a Deaf person can carry and use independently, without booking an "
     "interpreter and without surrendering the privacy of the exchange, restores a "
     "measure of self-determination. This consideration directly shaped two design "
     "decisions in Together: the privacy guarantee that no raw video leaves the "
     "device, and the offline fallback path that keeps the tool working without a "
     "dependence on a remote service that could log the conversation.")

h2("1.5 Problem Statement")
body("The central problem this thesis addresses can be stated as follows:")
para("How can we build an accessible, low-cost, real-time system that enables "
     "two-way communication between sign-language users and non-signers, in both "
     "American and Arabic sign languages, using only a standard web browser and "
     "commodity hardware, while respecting the distinct grammar of sign languages "
     "and remaining usable even without a constant cloud connection?",
     italic=True, space_before=4, space_after=8)
body("Answering this question requires solving several coupled sub-problems: "
     "robust, low-latency recognition of signs from camera-derived landmarks; "
     "linguistically faithful translation between gloss and natural language in two "
     "languages; synthesis of sign animations for arbitrary input including "
     "out-of-vocabulary words; real-time peer-to-peer media for live conversations; "
     "and a production-grade, secure, and resilient platform to host all of the "
     "above. Each of these sub-problems is treated in a dedicated chapter of this "
     "thesis.")

h2("1.6 Research Objectives")
numbered("To design and implement a bi-directional sign-language translation system "
         "supporting sign-to-text, sign-to-speech, text-to-sign, and speech-to-sign "
         "for both ASL/English and ArSL/Arabic.")
numbered("To engineer a browser-based recognition pipeline using MediaPipe Holistic "
         "landmarks, a TensorFlow-Lite ASL model, and a custom CNN–GRU Arabic model, "
         "with debouncing suited to continuous signing.")
numbered("To model sign-language grammar through bidirectional gloss translation "
         "(Topic–Comment ordering) and a non-manual-marker annotation layer.")
numbered("To build a text-to-sign avatar pipeline using Sentence-BERT semantic "
         "search over a pgvector index, with on-the-fly fingerspelling and motion "
         "stitching for out-of-vocabulary words.")
numbered("To implement a live two-party meeting mode over WebRTC with Socket.IO "
         "signalling and role-based interfaces.")
numbered("To guarantee resilience and affordability through a pluggable provider "
         "abstraction that falls back from cloud to fully offline language, speech, "
         "and synthesis services.")
numbered("To secure and productionise the platform with JWT authentication, "
         "Argon2id hashing, rate limiting, and an indexed PostgreSQL/pgvector data "
         "layer, and to optimise its end-to-end latency.")
numbered("To deliver a fully bilingual, accessible, themed user interface, and to "
         "evaluate the system both quantitatively and qualitatively.")

h2("1.7 Scope and Limitations")
h3("1.7.1 Scope")
bullet("Isolated-sign recognition for a 250-word ASL vocabulary and a 20-word "
       "Arabic vocabulary, debounced into running translation.")
bullet("Bidirectional gloss/sentence translation and a non-manual-marker hook.")
bullet("Text/speech-to-sign avatar playback with semantic sign lookup and stitching.")
bullet("A real-time meeting mode, authentication, a bilingual web UI, and a "
       "resilient provider layer with offline fallbacks.")
h3("1.7.2 Limitations")
bullet("The system recognises isolated signs assembled by a debouncing buffer "
       "rather than performing full continuous sign-language recognition with "
       "co-articulation modelling; this is a deliberate, well-scoped simplification.")
bullet("The Arabic vocabulary is limited to 20 high-frequency signs, constrained by "
       "the availability of curated training data.")
bullet("The non-manual-marker layer is currently a rule-based hook; full facial "
       "grammar rendering on the avatar is left as future work.")
bullet("Recognition accuracy depends on lighting, camera quality, and signer "
       "positioning, as with any vision-based system.")

h2("1.8 Thesis Organisation")
body("The remainder of this thesis is organised as follows.")
org = [
    ("Chapter 2 — Background and Literature Review", "surveys sign-language linguistics, the deep-learning foundations of sequence and gesture recognition, and existing commercial and academic translation systems, positioning Together against them."),
    ("Chapter 3 — Requirements Analysis", "defines stakeholders, functional and non-functional requirements, actors, use cases, and the system's behavioural models."),
    ("Chapter 4 — System Architecture and Design", "presents the overall architecture, component decomposition, technology stack, data model, and guiding design principles."),
    ("Chapter 5 — Landmark Extraction and Feature Engineering", "details MediaPipe Holistic, the landmark schema, normalisation, and the feature representation fed to the models."),
    ("Chapter 6 — Sign-to-Text Recognition Models", "describes the ASL TensorFlow-Lite model and the Arabic CNN–GRU model, their training, inference, and the debouncing buffer."),
    ("Chapter 7 — Gloss, Syntax and Non-Manual Markers", "covers Topic–Comment grammar, bidirectional gloss translation, and the non-manual-marker annotation layer."),
    ("Chapter 8 — Text and Speech to Sign", "explains the avatar pipeline: SBERT semantic lookup over pgvector and the gloss-and-stitch synthesis of out-of-vocabulary words."),
    ("Chapter 9 — The Provider Abstraction", "presents the pluggable LLM/TTS/STT layer and the cloud-to-offline fallback chains."),
    ("Chapter 10 — Real-Time Meeting Mode", "describes the WebRTC media path, Socket.IO signalling, and role-based interfaces."),
    ("Chapter 11 — Authentication, Security and Data Layer", "details Argon2id hashing, JWT access/refresh tokens, rate limiting, and the PostgreSQL/pgvector schema."),
    ("Chapter 12 — Frontend and Design System", "covers the bilingual, RTL-aware, themed design-token UI and the accessibility features."),
    ("Chapter 13 — Performance Engineering and Latency", "documents the profiling infrastructure and the optimisations that keep the system responsive."),
    ("Chapter 14 — Evaluation and Testing", "presents the experimental methodology, the model baseline results, and the automated test suite."),
    ("Chapter 15 — Conclusion and Future Work", "summarises the contributions and outlines directions for extending the system."),
]
for title, desc in org:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(title + " ").bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════ CHAPTER 2 ═══════════════════════════
h1("Chapter 2: Background and Literature Review")

h2("2.1 Introduction")
body("This chapter establishes the conceptual and technical foundations on which "
     "Together is built and situates the project within the existing landscape of "
     "sign-language technology. It is organised in three movements: first, the "
     "linguistics of sign languages, because faithful translation is impossible "
     "without understanding how sign languages actually work; second, the "
     "deep-learning building blocks for recognising motion sequences and producing "
     "language; and third, a critical survey of existing commercial and academic "
     "systems, culminating in a feature comparison that motivates the contributions "
     "of Together.")

h2("2.2 The Linguistics of Sign Languages")
h3("2.2.1 Sign Languages Are Natural Languages")
body("A recurring misconception is that sign language is a universal pantomime or a "
     "manual encoding of a spoken language. In fact there are hundreds of distinct "
     "sign languages — ASL, British Sign Language, Arabic Sign Language, and many "
     "regional variants — that are mutually unintelligible and grammatically "
     "independent of the surrounding spoken languages. ASL, for example, is "
     "historically related to French Sign Language rather than to English. Each "
     "sign language has a phonology (the sub-lexical parameters of handshape, "
     "location, movement, orientation, and non-manual features), a morphology, and "
     "a syntax of its own.")
h3("2.2.2 The Manual Parameters of a Sign")
body("A single sign is conventionally analysed along five simultaneous parameters: "
     "the handshape, the location at which the sign is made relative to the body, "
     "the movement trajectory, the palm orientation, and accompanying non-manual "
     "features. Two signs may differ in only one parameter and yet carry entirely "
     "different meanings, which is precisely why a recognition system must capture "
     "fine-grained, time-resolved hand and body configuration rather than coarse "
     "gestures. This directly justifies the use of MediaPipe Holistic, which "
     "provides 21 landmarks per hand in addition to body pose and facial mesh.")
h3("2.2.3 Topic–Comment Structure")
body("Whereas English is predominantly a Subject–Verb–Object (SVO) language, ASL "
     "frequently employs a Topic–Comment structure: the topic of the utterance — "
     "often the object, place, or time — is established first, and the comment "
     "about it follows. The sentence “I am going to the store” is naturally "
     "signed as STORE I GO, fronting the location. Articles (a, an, the) and "
     "copulas (is, are, am) are typically dropped because they carry no independent "
     "sign. Wh-question words are often placed at the end of the clause: "
     "“What is your name?” becomes YOUR NAME WHAT. Any system that maps "
     "signs to text word-for-word, ignoring this reordering, will produce output "
     "that is at best stilted and at worst unintelligible. Together explicitly "
     "models this transformation in both directions, as detailed in Chapter 7.")
table(
    ["English (SVO)", "ASL Gloss (Topic–Comment)"],
    [["I am going to the store.", "STORE I GO"],
     ["What is your name?", "YOUR NAME WHAT"],
     ["My mother loves me.", "MOTHER LOVE ME"],
     ["Do you want to eat?", "YOU WANT EAT YOU"],
     ["I have finished eating.", "FINISH EAT I"]],
    caption="Representative mappings between English sentences and ASL gloss, "
            "illustrating Topic–Comment reordering and the dropping of "
            "articles and copulas.")
h3("2.2.4 Non-Manual Markers")
body("A large part of sign-language grammar is carried not by the hands but by the "
     "face and body — the non-manual markers (NMM). Raised eyebrows mark a yes/no "
     "question; furrowed brows mark a wh-question; a headshake conveys negation; a "
     "brow-raise combined with a head tilt marks the topic of a Topic–Comment "
     "construction. These markers span several manual signs simultaneously rather "
     "than occupying their own slot in the sequence. A complete translation system "
     "must eventually model them; Together introduces a dedicated, documented "
     "non-manual-marker annotation layer (Section 7.4) that classifies sentence "
     "type and attaches the appropriate marker span, providing a clean hook for an "
     "avatar to render facial grammar.")

h2("2.3 Deep Learning Foundations")
body("This section reviews the neural-network components relevant to sign "
     "recognition and language generation, building up from the perception of "
     "spatial structure to the modelling of temporal sequences.")
h3("2.3.1 Convolutional Neural Networks")
body("Convolutional Neural Networks (CNNs) apply learnable filters across an input "
     "to detect local patterns, sharing weights across positions and thereby "
     "achieving translation invariance and parameter efficiency. While CNNs are most "
     "associated with two-dimensional images, one-dimensional convolutions are "
     "equally effective at extracting local temporal patterns from a sequence of "
     "feature vectors. In the Arabic model of Together, two stacked 1-D "
     "convolutional layers act as a learned, short-window feature extractor over the "
     "landmark time series before any recurrent processing, capturing local motion "
     "primitives such as the onset of a handshape change.")
figure_placeholder("A convolutional layer applying learnable filters across the "
                   "input to extract local features, with weight sharing across "
                   "positions.")
h3("2.3.2 Recurrent Networks, LSTM and GRU")
body("Sign language is intrinsically temporal: the same handshape can begin two "
     "different signs, and meaning emerges from the trajectory over time. Recurrent "
     "Neural Networks (RNNs) process a sequence step by step while maintaining a "
     "hidden state that summarises everything seen so far. Vanilla RNNs, however, "
     "struggle to learn long-range dependencies because gradients vanish or explode "
     "over many time steps. The Long Short-Term Memory (LSTM) cell solves this with "
     "an explicit memory cell and input, forget, and output gates that regulate the "
     "flow of information, allowing the network to retain relevant context across "
     "long sequences.")
body("The Gated Recurrent Unit (GRU) is a streamlined alternative that merges the "
     "forget and input gates into a single update gate and combines the cell and "
     "hidden states, yielding comparable accuracy with fewer parameters and faster "
     "training and inference. Because Together's Arabic model must run with low "
     "latency on commodity CPUs, the GRU's efficiency makes it the natural choice; "
     "the model uses a two-layer bidirectional GRU on top of the convolutional "
     "front-end.")
figure_placeholder("Comparison of a vanilla RNN cell and an LSTM cell, showing the "
                   "gating mechanism that mitigates the vanishing-gradient problem.")
h3("2.3.3 Bidirectional Recurrence")
body("In many recognition settings the full clip of a sign is available at once "
     "rather than strictly streaming, so the network may legitimately look both "
     "forward and backward in time. A bidirectional recurrent layer runs two "
     "passes, one in each temporal direction, and concatenates their hidden states, "
     "so the representation at each frame is informed by both past and future "
     "context. This is particularly valuable for disambiguating signs whose meaning "
     "only becomes clear once the movement completes. Together's Arabic model uses a "
     "bidirectional GRU; the prior generation of the project (documented in earlier "
     "work) used a bidirectional LSTM, and the migration to a CNN–GRU reflects a "
     "deliberate trade of a small amount of capacity for substantially lower "
     "latency.")
figure_placeholder("A bidirectional recurrent layer running forward and backward "
                   "passes and concatenating their hidden states at each time step.")
h3("2.3.4 Landmark-Based versus Pixel-Based Recognition")
body("Two broad strategies exist for vision-based sign recognition. Pixel-based "
     "approaches feed raw video frames (or 3-D convolutional stacks of frames) "
     "directly to a network, which must learn to localise the hands and body before "
     "it can recognise anything. Landmark-based approaches first extract a compact "
     "skeleton of keypoints — joint coordinates for the hands, body, and face — and "
     "operate on this low-dimensional representation. Together adopts the "
     "landmark-based strategy via MediaPipe Holistic for several decisive reasons: "
     "the representation is two to three orders of magnitude smaller than raw video, "
     "which makes browser-to-server streaming feasible; it is largely invariant to "
     "background, clothing, and lighting; it preserves user privacy because no raw "
     "imagery leaves the browser; and it lets a small, fast model achieve good "
     "accuracy. The trade-off is a dependence on the quality of the landmark "
     "extractor, which we mitigate with interpolation of dropped frames.")
h3("2.3.5 Transformers, Embeddings and Large Language Models")
body("The translation between gloss and natural language is a sequence-to-sequence "
     "language task, not a perception task, and modern Large Language Models (LLMs) "
     "built on the Transformer architecture excel at it. Rather than training a "
     "bespoke translation network, Together uses few-shot prompting of an LLM to "
     "convert gloss to sentence and sentence to gloss, which generalises far beyond "
     "any hand-written rule set while requiring no task-specific training data. For "
     "semantic sign lookup, the system uses Sentence-BERT (SBERT), which maps a word "
     "or phrase to a 384-dimensional embedding such that semantically similar terms "
     "lie close together; this allows a query like “mom” to retrieve the "
     "stored sign for “mother.” The embeddings are indexed with pgvector "
     "for fast approximate nearest-neighbour search.")

h3("2.3.6 Optimisation, Regularisation and Training Dynamics")
body("Training a recognition network well depends on more than its architecture. "
     "The Arabic model is trained with the cross-entropy loss appropriate to "
     "multi-class classification and regularised heavily — a dropout of 0.3 inside "
     "the recurrent layers and 0.5 before the final classifier — because the "
     "vocabulary is small and the risk of over-fitting to a limited number of signers "
     "is real. Batch normalisation after each convolution stabilises and accelerates "
     "training by keeping the distribution of activations well-conditioned across "
     "layers. Adaptive optimisers such as Adam and its weight-decayed variant AdamW "
     "are the natural choice for this kind of network because they adapt the "
     "effective learning rate per parameter from running estimates of the gradient's "
     "first and second moments, which makes training robust to the scale differences "
     "between the convolutional and recurrent parameters. Weight decay, applied "
     "correctly as in AdamW, further discourages large weights and improves "
     "generalisation.")
body("Two further techniques materially affect deployed accuracy. The first is the "
     "multi-window temporal ensembling described in Chapter 5, which is in effect a "
     "form of test-time augmentation: by averaging predictions over windows of "
     "different temporal extents the model becomes robust to variation in signing "
     "speed without any change to its weights. The second is the confidence "
     "threshold, which trades recall for precision; in an interactive system it is "
     "far better to emit nothing during an ambiguous transition than to emit a wrong "
     "sign that the downstream translator will faithfully turn into a wrong word.")

h2("2.4 Existing Sign-Language Translation Systems")
body("We now survey representative existing systems, with attention to their "
     "directionality, language coverage, hardware requirements, and grammatical "
     "fidelity.")
h3("2.4.1 Hand Talk")
body("Hand Talk is a widely cited mobile application that translates written and "
     "spoken Portuguese (and later English) into a signing 3-D avatar (“Hugo”). "
     "It is polished and popular for education, but it is fundamentally "
     "uni-directional — it produces signs from text/speech but does not recognise "
     "signs from a camera — and it does not address Arabic sign language. It "
     "demonstrates the viability and appeal of avatar-based output, which Together "
     "also provides, while leaving the harder recognition direction unaddressed.")
h3("2.4.2 SignAll")
body("SignAll is among the most technically ambitious commercial systems, "
     "performing continuous ASL-to-text recognition. Historically it relied on a "
     "multi-camera rig and, in some configurations, coloured gloves to achieve "
     "reliable hand tracking, which makes it accurate but expensive and impractical "
     "for spontaneous, everyday use on a personal device. It illustrates both the "
     "promise of genuine sign recognition and the cost of hardware-heavy approaches "
     "that Together explicitly avoids by using only a commodity webcam and a "
     "landmark extractor.")
h3("2.4.3 Google Sign-Language Detection")
body("Google has published work on real-time sign-language detection for video "
     "conferencing — determining when a participant is signing so they can be "
     "visually promoted as the active speaker. This is a valuable accessibility "
     "primitive, but detecting that someone is signing is a far simpler problem than "
     "recognising what they are signing and translating it. Together performs the "
     "full recognition-and-translation task and additionally builds an end-to-end "
     "meeting mode around it.")
h3("2.4.4 Other Notable Systems and Academic Work")
body("A substantial body of academic work tackles isolated and continuous sign "
     "recognition on datasets such as the large-scale word-level ASL benchmarks, "
     "using 3-D CNNs, recurrent networks, graph convolutional networks over "
     "skeletons, and Transformer encoders. These works push recognition accuracy on "
     "fixed benchmarks but are typically research prototypes rather than deployed, "
     "bidirectional, multilingual, production systems with authentication, a meeting "
     "mode, and offline fallbacks. A handful of dictionary-style apps provide "
     "word-to-sign video lookups but no recognition and no grammar. The gap that "
     "Together fills is integration: combining recognition, grammar-aware "
     "translation, synthesis, real-time communication, and resilience into one "
     "deployable, bilingual application.")

h2("2.5 Comparison with Together")
table(
    ["Capability", "Hand Talk", "SignAll", "Google Detect", "Dictionary apps", "Together"],
    [
        ["Sign → text", "No", "Yes", "Detect only", "No", "Yes"],
        ["Text/speech → sign avatar", "Yes", "No", "No", "Lookup only", "Yes"],
        ["Bi-directional", "No", "No", "No", "No", "Yes"],
        ["Arabic sign language", "No", "No", "No", "Rare", "Yes"],
        ["Commodity hardware only", "Yes", "No", "Yes", "Yes", "Yes"],
        ["Grammar-aware (Topic–Comment)", "Partial", "Partial", "N/A", "No", "Yes"],
        ["Live two-party meeting", "No", "No", "Partial", "No", "Yes"],
        ["Works offline (fallbacks)", "No", "No", "No", "Partial", "Yes"],
        ["Install-free (browser)", "App", "No", "N/A", "App", "Yes"],
    ],
    caption="Feature comparison of representative sign-language systems against "
            "Together. Together is distinguished by being simultaneously "
            "bi-directional, bilingual, grammar-aware, hardware-light, and "
            "resilient to network loss.")
figure_placeholder("Feature comparison chart of sign-language translation systems.")

h2("2.6 Contributions of Together")
body("Against this backdrop, the specific contributions of this project are:")
numbered("A single, install-free web application that delivers all four translation "
         "flows (sign↔text and sign↔speech) for two languages.")
numbered("A landmark-based recognition pipeline combining a 250-class ASL "
         "TensorFlow-Lite model and a purpose-built CNN–GRU Arabic model, with a "
         "debouncing buffer tuned for continuous signing.")
numbered("Grammar-aware, bidirectional gloss translation through few-shot LLM "
         "prompting, plus a documented non-manual-marker annotation layer.")
numbered("A text-to-sign synthesis pipeline with SBERT semantic lookup over an "
         "indexed pgvector store and a Savitzky–Golay-smoothed fingerspelling "
         "stitcher for out-of-vocabulary words.")
numbered("A resilient provider abstraction that degrades gracefully from cloud to "
         "fully offline language, speech, and synthesis services.")
numbered("A production-grade platform: JWT/Argon2id security, an indexed "
         "PostgreSQL/pgvector data layer, a real-time WebRTC meeting mode, a "
         "bilingual RTL-aware design system, and an end-to-end latency-optimisation "
         "phase with live profiling.")

# ═══════════════════════════ CHAPTER 3 ═══════════════════════════
h1("Chapter 3: Analysis and Requirements")

h2("3.1 Introduction to Problem Analysis")
body("Before designing a system it is essential to analyse precisely what it must "
     "do, for whom, and under what constraints. This chapter applies a standard "
     "requirements-engineering process to Together. It identifies the stakeholders "
     "and actors, enumerates functional and non-functional requirements, derives use "
     "cases, and models the system's behaviour and data through use-case, sequence, "
     "class, and entity-relationship diagrams. These artefacts form the contract "
     "that the architecture and implementation chapters fulfil.")

h2("3.2 Stakeholders")
body("A stakeholder is any party with a legitimate interest in the system. For "
     "Together we identify the following.")
bullet("Deaf and Hard-of-Hearing users: the primary beneficiaries, who sign to be "
       "understood by non-signers and who consume sign output from spoken input.")
bullet("Hearing, non-signing users: who speak or type and wish to communicate with "
       "a signer; in meeting mode they are the “Speaker” role.")
bullet("Educators and learners of sign language: who use the word-lookup and avatar "
       "features as a study aid.")
bullet("Institutions (hospitals, banks, schools, government services): potential "
       "deployers who need an accessible communication channel at their counters.")
bullet("Developers and maintainers: who extend the vocabulary, swap providers, and "
       "operate the platform.")
bullet("The project supervisors and examiners: who assess the academic and "
       "engineering quality of the work.")
figure_placeholder("Stakeholder map for the Together system.")

h2("3.3 Actors and Goals")
table(
    ["Actor", "Primary goals"],
    [
        ["Signer (DHH user)", "Be understood by non-signers; read/see translated speech as sign or caption"],
        ["Speaker (hearing user)", "Understand a signer; have their speech rendered as sign to the signer"],
        ["Registered user", "Authenticate, manage a profile, access dashboards and meeting rooms"],
        ["Learner", "Look up individual signs and watch the avatar reproduce them"],
        ["System administrator", "Operate the service, monitor latency metrics, manage providers"],
    ],
    caption="Actors of the Together system and their primary goals.")

h2("3.4 Functional Requirements")
body("The functional requirements specify the behaviours the system must exhibit. "
     "They are grouped by subsystem and given identifiers for traceability.")
fr = [
    ("FR-1", "The system shall extract pose, hand, and face landmarks from the user's webcam in the browser using MediaPipe Holistic."),
    ("FR-2", "The system shall recognise ASL signs from streamed landmark sequences using a 250-class model and return a gloss with a confidence score."),
    ("FR-3", "The system shall recognise Arabic signs from landmark sequences using a dedicated 20-class model."),
    ("FR-4", "The system shall debounce recognised signs using a vote-and-cooldown buffer to suppress duplicates and spurious detections."),
    ("FR-5", "The system shall translate a sequence of recognised glosses into a natural sentence in the appropriate spoken language, honouring Topic–Comment grammar."),
    ("FR-6", "The system shall synthesise speech audio from a translated sentence (sign-to-speech)."),
    ("FR-7", "The system shall convert a natural sentence into ordered gloss tokens annotated with non-manual markers (text-to-sign)."),
    ("FR-8", "The system shall transcribe spoken audio to text (speech-to-sign input)."),
    ("FR-9", "The system shall map each gloss token to a stored landmark sequence via exact and semantic (SBERT) matching, and play it back on an avatar."),
    ("FR-10", "The system shall synthesise animations for out-of-vocabulary words by fingerspelling and stitching, or report the word as unrenderable."),
    ("FR-11", "The system shall provide a live two-party meeting mode connecting a signer and a speaker over WebRTC with per-role interfaces."),
    ("FR-12", "The system shall allow users to register, log in, refresh sessions, and log out securely."),
    ("FR-13", "The system shall provide a bilingual interface (English LTR and Arabic RTL) with a light/dark theme."),
    ("FR-14", "The system shall expose health and latency-metrics endpoints for monitoring."),
    ("FR-15", "The system shall fall back from cloud to offline providers for language, speech synthesis, and transcription when necessary."),
]
table(["ID", "Functional Requirement"], fr)

h2("3.5 Non-Functional Requirements")
nfr = [
    ("NFR-1 Performance", "End-to-end translation latency shall be low enough to feel conversational; blocking inference must not stall unrelated requests."),
    ("NFR-2 Resilience", "The core experience shall continue to function when cloud providers are unavailable, via offline fallbacks."),
    ("NFR-3 Security", "Passwords shall be hashed with Argon2id; sessions shall use signed JWTs with rotating, revocable refresh tokens; auth endpoints shall be rate-limited."),
    ("NFR-4 Privacy", "Raw camera imagery shall never leave the browser; only derived landmarks are transmitted."),
    ("NFR-5 Accessibility", "The UI shall support keyboard navigation, focus-visible rings, reduced-motion preferences, and full RTL mirroring for Arabic."),
    ("NFR-6 Scalability", "Semantic sign search shall use an indexed vector store rather than in-memory brute force, to scale with vocabulary size."),
    ("NFR-7 Maintainability", "External services shall be accessed only through abstract interfaces so providers can be swapped without touching business logic."),
    ("NFR-8 Portability", "The system shall run via Docker Compose and require only commodity hardware and a standard browser."),
    ("NFR-9 Observability", "Per-stage timing shall be measurable live through a metrics endpoint."),
    ("NFR-10 Usability", "The interface shall be learnable without instruction for both signer and speaker roles."),
]
table(["Category", "Non-Functional Requirement"], nfr)

h2("3.6 Use Cases")
body("The principal use cases realise the functional requirements as user-facing "
     "interactions. Three representative use cases are specified in detail below; "
     "the full set is captured in the use-case diagram.")
h3("3.6.1 Use Case: Translate Sign to Speech")
table(
    ["Field", "Description"],
    [
        ["Name", "Translate Sign to Speech"],
        ["Actor", "Signer (DHH user)"],
        ["Precondition", "User is authenticated and has granted camera access"],
        ["Main flow", "1. User selects the sign-to-speech module and language. 2. Browser extracts landmarks and streams them. 3. Server recognises signs and debounces them into glosses. 4. The gloss sequence is translated into a sentence. 5. The sentence is synthesised to speech and played."],
        ["Alternative flow", "If no sign is confidently detected, no gloss is emitted and the buffer continues."],
        ["Postcondition", "Spoken audio of the translated sentence is produced"],
    ],
    caption="Use-case specification: Translate Sign to Speech.")
h3("3.6.2 Use Case: Translate Speech to Sign")
table(
    ["Field", "Description"],
    [
        ["Name", "Translate Speech to Sign"],
        ["Actor", "Speaker (hearing user)"],
        ["Precondition", "User is authenticated and has granted microphone access"],
        ["Main flow", "1. User speaks. 2. Audio is transcribed to text. 3. Text is converted to gloss tokens with NMM annotations. 4. Each token is matched to a stored sign sequence. 5. The avatar plays the signs in order."],
        ["Alternative flow", "Out-of-vocabulary words are fingerspelled via stitching, or reported as unrenderable."],
        ["Postcondition", "The avatar has signed the spoken utterance"],
    ],
    caption="Use-case specification: Translate Speech to Sign.")
h3("3.6.3 Use Case: Conduct a Live Meeting")
table(
    ["Field", "Description"],
    [
        ["Name", "Conduct a Live Meeting"],
        ["Actor", "Signer and Speaker"],
        ["Precondition", "Both users are authenticated and join the same room"],
        ["Main flow", "1. Users join a room and announce presence. 2. WebRTC offer/answer and ICE candidates are exchanged via Socket.IO. 3. A peer-to-peer media connection is established. 4. Each side sees captions or an avatar appropriate to its role while translation runs continuously."],
        ["Alternative flow", "On connection loss the signalling layer attempts re-negotiation."],
        ["Postcondition", "A two-way translated conversation has taken place"],
    ],
    caption="Use-case specification: Conduct a Live Meeting.")
figure_placeholder("Use-case diagram of the Together system, showing the Signer, "
                   "Speaker, Learner, and Administrator actors and their use cases.")

h2("3.7 System Sequence Diagrams")
body("Sequence diagrams capture the temporal ordering of messages between the user, "
     "the browser, the server, the models, and the providers for each principal "
     "flow. For the sign-to-speech flow, the browser repeatedly sends landmark "
     "frames to the recognition endpoint, which returns glosses; once the "
     "debouncing buffer accepts a stable gloss sequence, the client requests a "
     "sentence translation, then requests speech synthesis. For the speech-to-sign "
     "flow the ordering is mirrored: transcription, then gloss conversion, then "
     "batched sign lookup, then avatar playback.")
figure_placeholder("System sequence diagram for the sign-to-speech translation flow.")
figure_placeholder("System sequence diagram for the speech-to-sign translation flow.")

h2("3.8 Class Diagram and Interface Specification")
body("At the design level the server decomposes into cohesive classes and modules. "
     "The SignDB class encapsulates English sign lookup backed by Postgres and "
     "pgvector; ArabicSignDB specialises it for the Arabic vocabulary; and "
     "ASLService extends SignDB with the TensorFlow-Lite gesture recogniser. The "
     "provider layer is defined by abstract base classes — LLMProvider, TTSProvider, "
     "and STTProvider — each with a single primary method (generate, synthesize, "
     "transcribe) and an availability check, against which concrete Gemini, Ollama, "
     "and local adapters are written. The data layer is modelled by ORM classes "
     "User, RefreshToken, and Sign. This separation of interface from implementation "
     "is what makes the provider chains and the storage backend swappable.")
figure_placeholder("Class diagram of the Together server, showing SignDB / "
                   "ArabicSignDB / ASLService, the provider interfaces, and the ORM "
                   "entities.")

h2("3.9 Entity–Relationship Diagram")
body("The persistent data model comprises three principal entities. A User has "
     "credentials and a role; a User owns zero or more RefreshTokens that implement "
     "session rotation and revocation; and a Sign records the metadata and SBERT "
     "embedding for one vocabulary item in one language, with its raw landmark "
     "sequence stored on disk and referenced by path. The Sign entity carries a "
     "unique constraint on the (word, language) pair and an HNSW vector index on the "
     "embedding column to support fast semantic search.")
figure_placeholder("Entity–Relationship diagram of the Together database (User, "
                   "RefreshToken, Sign).")

# ═══════════════════════════ CHAPTER 4 ═══════════════════════════
h1("Chapter 4: System Architecture and Design")

h2("4.1 System Overview")
body("Together is a client–server web application. The client is a server-rendered, "
     "framework-free front end that runs MediaPipe Holistic, captures audio, drives "
     "the WebRTC peer connection for meetings, and renders the avatar. The server is "
     "a FastAPI application, augmented with a Socket.IO server for real-time "
     "signalling, that hosts the recognition models, the gloss and synthesis logic, "
     "the provider chains, authentication, and the data layer. Persistent state "
     "lives in PostgreSQL with the pgvector extension, and raw landmark sequences "
     "live as compressed files on disk. The whole system is reproducible through "
     "Docker Compose, which brings up the database, the web application, and a local "
     "Ollama instance for the offline LLM fallback.")
figure_placeholder("High-level system architecture diagram of Together: browser "
                   "client, FastAPI + Socket.IO server, models, provider chains, and "
                   "the PostgreSQL/pgvector + on-disk landmark store.")

h2("4.2 Architectural Principles")
body("Four deliberate decisions anchor the architecture, and every subsequent "
     "chapter can be read as a consequence of one of them.")
numbered("The front end is server-rendered HTML, CSS, and vanilla JavaScript with "
         "Jinja2 templates and a design-token system — no single-page-application "
         "framework and no build step — which keeps the client simple, fast to load, "
         "and trivial to theme and mirror for RTL.")
numbered("The machine-learning stack is split by language and optimised for CPU: "
         "TensorFlow Lite (via ai-edge-litert) for the 250-class English model and a "
         "PyTorch CNN–GRU for the 20-class Arabic model, with MediaPipe Holistic "
         "extracting landmarks in the browser.")
numbered("Real-time communication uses Socket.IO for WebRTC signalling so that "
         "meeting media flows peer-to-peer while only small signalling messages "
         "traverse the server.")
numbered("Both languages are first-class, and every external dependency is accessed "
         "through a pluggable provider chain that degrades from cloud to offline, so "
         "the system is simultaneously bilingual and resilient.")

h2("4.3 Component Decomposition")
h3("4.3.1 The Recognition Subsystem")
body("The recognition subsystem owns the path from landmarks to glosses. It is "
     "embodied in asl_service.py, which defines SignDB and its specialisations and "
     "the ASLService that wraps the TensorFlow-Lite interpreter, and in "
     "sign_predictor.py, which defines the PyTorch CNN–GRU model and its "
     "preprocessing. Both expose a simple predict interface that accepts a landmark "
     "array and returns a label and confidence, decoupling the rest of the server "
     "from the details of either ML framework.")
h3("4.3.2 The Language Subsystem")
body("The language subsystem owns the path between glosses and natural language. It "
     "comprises gloss.py (bidirectional gloss translation and non-manual-marker "
     "annotation) and the provider layer that supplies the LLM. Because translation "
     "is delegated to an LLM through a narrow interface, the same code serves both "
     "English and Arabic by swapping a few-shot exemplar set.")
h3("4.3.3 The Synthesis Subsystem")
body("The synthesis subsystem owns the path from glosses to renderable motion. It "
     "comprises the sign-lookup methods of SignDB (exact, contraction, and semantic "
     "matching), the stitch.py module that fingerspells and smooths out-of-vocabulary "
     "words, and the landmark_store that loads motion sequences from disk on demand.")
h3("4.3.4 The Real-Time Subsystem")
body("The real-time subsystem owns the meeting mode. It comprises the Socket.IO "
     "event handlers for room management and WebRTC signalling and the client-side "
     "peer-connection logic. Media never passes through the server; only offers, "
     "answers, and ICE candidates do.")
h3("4.3.5 The Platform Subsystem")
body("The platform subsystem owns cross-cutting concerns: authentication and "
     "security (auth.py), the database and repository layer (db/), request "
     "profiling (profiling.py), configuration, and the HTTP/ASGI plumbing in "
     "main.py.")

h2("4.4 Technology Stack")
table(
    ["Layer", "Technology", "Role"],
    [
        ["Web framework", "FastAPI + Uvicorn (ASGI)", "HTTP API, routing, async request handling"],
        ["Real-time", "python-socketio", "WebRTC signalling, room presence"],
        ["Templating", "Jinja2", "Server-rendered bilingual pages"],
        ["ASL model", "TensorFlow Lite (ai-edge-litert)", "250-class sign recognition with XNNPACK"],
        ["Arabic model", "PyTorch (CNN–GRU)", "20-class sign recognition, INT8 on CPU"],
        ["Landmarks", "MediaPipe Holistic", "Pose/hand/face keypoint extraction in browser"],
        ["Semantic search", "Sentence-BERT (all-MiniLM-L6-v2)", "384-d embeddings for sign lookup"],
        ["Database", "PostgreSQL + pgvector", "Users, tokens, sign metadata, HNSW vector index"],
        ["Migrations", "Alembic", "Schema versioning"],
        ["Auth", "Argon2id, python-jose (JWT), bcrypt (legacy)", "Hashing and token security"],
        ["LLM providers", "Google Gemini → Ollama", "Gloss translation"],
        ["TTS providers", "Gemini TTS → pyttsx3", "Speech synthesis"],
        ["STT providers", "Gemini → faster-whisper", "Speech transcription"],
        ["Signal processing", "SciPy (Savitzky–Golay)", "Stitch smoothing"],
        ["Packaging", "Docker Compose", "Reproducible multi-service deployment"],
    ],
    caption="The Together technology stack by layer.")

h2("4.5 Data Architecture")
body("A central design decision is the separation of small, queryable metadata from "
     "large, bulky motion data. Sign metadata and SBERT embeddings live in "
     "PostgreSQL, where the embedding is a pgvector column with an HNSW cosine index, "
     "so similarity search is a single indexed query rather than an in-Python "
     "brute-force loop, and no embeddings are held in RAM. The raw MediaPipe "
     "landmark sequences — arrays shaped [n_frames, 543, 3] — are far too large to "
     "store as rows, so they are kept as compressed .npz files on disk, keyed by "
     "sign id, and the database holds only the file path. Landmarks are then loaded "
     "lazily, only for the specific signs a request actually needs.")

h2("4.6 Design Principles")
bullet("Separation of concerns: each subsystem has a single responsibility and a "
       "narrow interface.")
bullet("Graceful degradation: every external dependency has a fallback, and every "
       "pipeline returns something renderable even when a stage fails.")
bullet("Privacy by design: only derived landmarks, never raw video, leave the "
       "browser.")
bullet("Observability: hot stages are timed and exposed so regressions are visible "
       "without external tooling.")
bullet("Internationalisation by construction: the UI uses CSS logical properties so "
       "Arabic RTL mirroring costs nothing, and both languages share one codebase.")

# ═══════════════════════════ CHAPTER 5 ═══════════════════════════
h1("Chapter 5: Landmark Extraction and Feature Engineering")

h2("5.1 Introduction")
body("Every recognition pipeline in Together begins not with pixels but with "
     "landmarks. This chapter describes how raw camera frames are reduced to a "
     "compact, privacy-preserving skeleton of keypoints by MediaPipe Holistic, how "
     "those keypoints are structured, and how they are normalised and assembled into "
     "the feature vectors consumed by the recognition models. The quality of this "
     "front end fundamentally bounds the accuracy of everything downstream, so it "
     "warrants careful treatment.")

h2("5.2 MediaPipe Holistic")
body("MediaPipe Holistic is a cross-platform machine-learning pipeline that "
     "simultaneously estimates human pose, hand, and face landmarks from a single "
     "RGB video frame. It runs in the browser via WebAssembly, which means the "
     "computationally intensive perception step happens entirely on the user's "
     "device; only the resulting numeric landmarks are streamed to the server. "
     "Holistic produces a total of 543 landmarks per frame: 33 pose landmarks, 468 "
     "face-mesh landmarks, and 21 landmarks for each of the two hands. Each landmark "
     "is a triple of normalised (x, y, z) coordinates, giving the canonical array "
     "shape of [n_frames, 543, 3] used throughout the system.")
figure_placeholder("MediaPipe Holistic keypoints: 33 pose, 468 face, and 21 per "
                   "hand, totalling 543 landmarks per frame.")
table(
    ["Group", "Landmark count", "Used by ASL model", "Used by Arabic model"],
    [
        ["Pose", "33 (17 upper-body used)", "Yes (full holistic array)", "Yes (17 upper-body joints)"],
        ["Left hand", "21", "Yes", "Yes"],
        ["Right hand", "21", "Yes", "Yes"],
        ["Face mesh", "468", "Yes (present in array)", "No (dropped for compactness)"],
        ["Total", "543", "543×3 input", "59 keypoints → 177 features"],
    ],
    caption="MediaPipe Holistic landmark groups and how each recognition model "
            "consumes them.")

h2("5.3 Why Landmarks Rather Than Pixels")
body("As argued in Chapter 2, the landmark representation is decisive for a "
     "browser-based, privacy-respecting, low-latency system. A single 543×3 float "
     "frame is a few kilobytes, whereas a raw camera frame is orders of magnitude "
     "larger; streaming landmarks rather than video makes real-time client–server "
     "operation feasible over ordinary connections. Because the representation "
     "abstracts away appearance, the model is insensitive to background clutter, "
     "skin tone, clothing, and lighting, improving fairness and robustness. And "
     "because raw imagery never leaves the device, the system sidesteps a large "
     "class of privacy concerns inherent in sending video of a person's face to a "
     "server.")

h2("5.4 The Optimised 59-Landmark Payload")
body("For the Arabic pipeline the client can transmit an optimised payload of only "
     "the 59 landmarks the model actually needs — 21 for the left hand, 17 for the "
     "upper-body pose, and 21 for the right hand — rather than the full 543. The "
     "server reconstructs the canonical 543-landmark array, slotting the received "
     "keypoints into their correct indices, before preprocessing. This reduces "
     "bandwidth roughly tenfold for the Arabic flow without any loss of information "
     "relevant to recognition, because the dropped face-mesh and lower-body "
     "landmarks are not used by that model.")
code_block([
    "# Reconstruct the 543-landmark array from a 59-landmark payload",
    "full = np.zeros((N, 543, 3), dtype=np.float32)",
    "full[:, 468:489, :] = frames[:, :21, :]    # Left hand  (21)",
    "full[:, 489:506, :] = frames[:, 21:38, :]   # Pose upper (17)",
    "full[:, 522:543, :] = frames[:, 38:59, :]   # Right hand (21)",
])

h2("5.5 Normalisation and Feature Construction")
body("Raw landmark coordinates are not directly comparable across users or "
     "distances from the camera: a sign made by a tall person standing close to the "
     "lens yields very different absolute coordinates from the same sign made by a "
     "shorter person standing farther away. The Arabic model therefore applies a "
     "carefully designed normalisation. First, the landscape camera coordinates are "
     "corrected to a 9:16 portrait virtual canvas of 720×1280 so that aspect ratio "
     "does not distort the geometry. Second, the depth (z) coordinate is zeroed, "
     "making recognition invariant to the signer's distance from the camera and "
     "avoiding reliance on the noisiest landmark dimension. Third, and most "
     "importantly, every coordinate is re-expressed relative to the centre point "
     "between the two shoulders and scaled by the shoulder width.")
body("This shoulder-relative normalisation is the key to user invariance: by "
     "subtracting the shoulder midpoint and dividing by the inter-shoulder distance, "
     "the representation becomes independent of where the signer is in the frame and "
     "of their absolute body size, capturing instead the configuration of the hands "
     "relative to the torso. After normalisation the model's feature vector for the "
     "Arabic pipeline concatenates 17 pose joints, 21 left-hand, and 21 right-hand "
     "landmarks — 59 keypoints — flattened to 177 features per frame.")
code_block([
    "center = (left_shoulder + right_shoulder) / 2.0",
    "scale  = ||left_shoulder - right_shoulder||   # inter-shoulder distance",
    "normalised = (landmark - center) / scale       # user- & distance-invariant",
])

h2("5.6 Handling Missing and Dropped Landmarks")
body("Hand tracking is imperfect: hands leave the frame, occlude one another, or "
     "move too fast for the tracker, producing frames with missing landmarks "
     "(zeros, or NaN from the holistic estimator). Naïvely feeding these gaps to a "
     "model degrades accuracy and produces visible jitter. Together addresses this "
     "with linear interpolation of missing hand coordinates: for any frame in which "
     "a hand is absent, the system finds the nearest visible frames before and after "
     "and linearly interpolates the hand position between them, falling back to the "
     "nearest available frame at the sequence boundaries. This reconstructs a smooth, "
     "plausible trajectory through short tracking dropouts. In the synthesis "
     "direction (Chapter 8), NaN landmarks are deliberately preserved rather than "
     "interpolated during smoothing, because smoothing a window containing NaN would "
     "corrupt good neighbouring frames.")

h2("5.7 Temporal Windowing")
body("Signs unfold over time, so recognition operates over a window of recent "
     "frames rather than a single frame. The system maintains a sliding history "
     "buffer of the most recent frames. To make recognition robust to the natural "
     "variation in how quickly different people sign, the Arabic pipeline performs "
     "multi-window inference: it samples three windows of the most recent 30, 45, and "
     "60 frames, uniformly resamples each to a fixed length of 30 frames, runs the "
     "model on all three, and averages the resulting class probabilities. This "
     "ensemble over temporal scales markedly stabilises predictions compared to a "
     "single fixed window. The ASL pipeline operates over a sliding window trimmed "
     "to a sequence length of 60 frames, padding shorter sequences by repeating the "
     "last frame.")
figure_placeholder("Temporal windowing and multi-scale sampling: the 30/45/60-frame "
                   "windows are each resampled to 30 frames and ensembled.")

# ═══════════════════════════ CHAPTER 6 ═══════════════════════════
h1("Chapter 6: Sign-to-Text Recognition Models")

h2("6.1 Introduction")
body("This chapter presents the two recognition models at the heart of the "
     "sign-to-text direction: the 250-class English ASL model deployed as a "
     "TensorFlow-Lite graph, and the purpose-built 20-class Arabic CNN–GRU model "
     "implemented in PyTorch. It explains each model's architecture, its inference "
     "procedure, the confidence thresholds that govern acceptance, and the "
     "debouncing buffer that turns a stream of per-window predictions into a clean "
     "sequence of glosses suitable for translation.")

h2("6.2 The English ASL Model (TensorFlow Lite)")
h3("6.2.1 Task and Vocabulary")
body("The English model recognises 250 distinct ASL signs — a vocabulary drawn from "
     "common everyday words — from a sequence of holistic landmarks. The label set "
     "is stored as an index-to-sign map, and the model is distributed as a "
     "TensorFlow-Lite (.tflite) file, a format optimised for efficient on-device "
     "and server-side CPU inference.")
h3("6.2.2 Why TensorFlow Lite")
body("Deploying the recognition model as a TensorFlow-Lite graph rather than a "
     "full TensorFlow model yields a small binary and fast CPU execution. The "
     "system loads the model through the ai-edge-litert interpreter, configured with "
     "multiple threads (the minimum of four and the available core count, overridable "
     "via an environment variable). LiteRT applies the XNNPACK delegate by default "
     "for floating-point CPU models, which accelerates the convolution and dense "
     "operations that dominate inference. The implementation gracefully handles "
     "older interpreter signatures that lack a thread-count argument.")
h3("6.2.3 Inference Procedure")
body("Inference accepts a landmark array of up to 60 frames. Sequences shorter than "
     "the fixed sequence length of 60 are padded by repeating the final frame; "
     "longer sequences are trimmed to the most recent 60. The batch is cast to the "
     "interpreter's expected dtype, the input tensor is resized and reallocated, the "
     "interpreter is invoked, and the output logits are read back. The class with "
     "the maximum activation is selected, and a softmax confidence is computed. A "
     "prediction is accepted only if the top logit exceeds a confidence threshold of "
     "0.80; otherwise the system reports that no sign was detected, which is the "
     "correct behaviour during the transitions between signs.")
code_block([
    "idx = argmax(logits)",
    "if logits[idx] > 0.80:                 # confidence gate",
    "    label = labels[idx]",
    "    softmax_conf = softmax(logits)[idx]",
    "    return label, softmax_conf",
    "return None, 0.0                       # transition / no confident sign",
])

h2("6.3 The Arabic Model (CNN–GRU)")
h3("6.3.1 Architecture")
body("The Arabic model, SignLanguageCNNGRU, is a compact hybrid network designed "
     "for low-latency CPU inference on the 20-word Arabic vocabulary. It takes a "
     "sequence of 177-dimensional feature vectors (the 59 normalised keypoints of "
     "Section 5.5) and processes them in two stages. A convolutional front end of "
     "two 1-D convolution layers, each with 128 channels, a kernel size of 3, batch "
     "normalisation, and ReLU activation, extracts local temporal motion features. A "
     "recurrent back end of a two-layer bidirectional GRU with 64 hidden units per "
     "direction then models the longer-range temporal structure of the sign. The "
     "final hidden state feeds a fully connected layer of 64 units with dropout and "
     "a final linear classifier over the 20 classes.")
table(
    ["Layer", "Configuration", "Output"],
    [
        ["Input", "30 frames × 177 features", "(B, 30, 177)"],
        ["Conv1d + BN + ReLU", "177→128, kernel 3, pad 1", "(B, 128, 30)"],
        ["Conv1d + BN + ReLU", "128→128, kernel 3, pad 1", "(B, 128, 30)"],
        ["Bidirectional GRU ×2", "128→64 per direction, dropout 0.3", "(B, 30, 128)"],
        ["Take last time step", "—", "(B, 128)"],
        ["Linear + ReLU", "128→64", "(B, 64)"],
        ["Dropout", "p = 0.5", "(B, 64)"],
        ["Linear (classifier)", "64→20", "(B, 20)"],
    ],
    caption="The SignLanguageCNNGRU architecture for Arabic sign recognition.")
figure_placeholder("Architecture of the CNN–GRU Arabic sign-recognition model.")
h3("6.3.2 INT8 Dynamic Quantization")
body("To minimise latency on commodity CPUs, the model applies INT8 dynamic "
     "quantization to its Linear and GRU weights at load time. Dynamic quantization "
     "stores weights as 8-bit integers and quantises activations on the fly during "
     "inference, shrinking the model and accelerating matrix multiplications with "
     "negligible accuracy impact for a network of this size. The optimisation is "
     "enabled by default on CPU and can be disabled via an environment variable for "
     "comparison. As reported in Chapter 13, this yields roughly a 1.8× speed-up "
     "over the FP32 baseline.")
h3("6.3.3 Multi-Window Inference and Top-3 Output")
body("As described in Section 5.7, the Arabic model performs multi-window inference, "
     "running the network on 30-, 45-, and 60-frame windows and averaging their "
     "softmax probabilities. The averaged distribution yields the predicted class, "
     "its confidence, and a top-3 list useful for interface feedback. A prediction "
     "is accepted only above a confidence threshold of 0.65, and a hand-visibility "
     "history is used to reset the buffer when the hands disappear and reappear, "
     "preventing stale frames from contaminating a new sign.")

h2("6.4 From Predictions to Glosses: The Debouncing Buffer")
body("A recognition model fired on every incoming window will emit many repeated and "
     "occasionally spurious predictions as a single sign is held and as the hands "
     "transition between signs. Feeding this raw stream to the translator would "
     "produce garbled, repetitive glosses. Together therefore interposes a "
     "vote-and-cooldown debouncing buffer between recognition and translation. The "
     "buffer accumulates recent predictions, accepts a gloss only when it has been "
     "voted for consistently, and then enters a cooldown during which the same sign "
     "will not be emitted again, so that holding a sign yields exactly one gloss. "
     "Confidence gating (the per-model thresholds above) and the hand-visibility "
     "reset complement the buffer, together producing a clean, de-duplicated gloss "
     "sequence that faithfully represents what was signed.")

h2("6.5 Comparison of the Two Models")
body("The two recognition models embody two reasonable but different points in the "
     "design space, and comparing them clarifies the engineering trade-offs. The ASL "
     "model is a larger, frozen TensorFlow-Lite graph covering 250 classes and "
     "consuming the full 543-landmark holistic array; it is deployed rather than "
     "trained within this project, and accelerated through XNNPACK and "
     "multi-threading. The Arabic model is a small, transparent PyTorch network "
     "covering 20 classes and consuming only the 59 keypoints that matter, trained "
     "and quantised as part of this work, and accelerated through INT8 dynamic "
     "quantization. The ASL model prioritises vocabulary breadth; the Arabic model "
     "prioritises latency and controllability on a focused vocabulary. Together they "
     "demonstrate that the surrounding pipeline — landmark extraction, debouncing, "
     "gloss translation, synthesis — is genuinely model-agnostic, depending only on "
     "the narrow predict interface each model exposes.")
table(
    ["Aspect", "ASL model", "Arabic model"],
    [
        ["Framework", "TensorFlow Lite (LiteRT)", "PyTorch"],
        ["Classes", "250", "20"],
        ["Input", "543 landmarks × 3", "59 keypoints → 177 features"],
        ["Architecture", "Pre-trained TFLite graph", "CNN (2×Conv1d) + 2-layer BiGRU"],
        ["Window", "60 frames (pad/trim)", "30/45/60 multi-window ensemble"],
        ["Confidence threshold", "0.80 (logit)", "0.65 (mean softmax)"],
        ["CPU acceleration", "XNNPACK + multithread", "INT8 dynamic quantization"],
    ],
    caption="Side-by-side comparison of the English ASL and Arabic recognition "
            "models.")

h2("6.6 Threading and the Event Loop")
body("Both models perform CPU-bound work that, if executed directly in the "
     "asynchronous event loop, would block every other in-flight request — including "
     "health checks and the meeting socket — until inference completed. All "
     "inference is therefore dispatched to a thread pool via Starlette's "
     "run_in_threadpool, so a single slow inference never serialises the server. "
     "This is examined quantitatively in Chapter 13.")

# ═══════════════════════════ CHAPTER 7 ═══════════════════════════
h1("Chapter 7: Gloss, Syntax and Non-Manual Markers")

h2("7.1 Introduction")
body("Recognising signs yields glosses; producing or consuming natural language "
     "requires translating between gloss order and spoken-language order. This "
     "chapter describes the bidirectional gloss-translation engine and the "
     "non-manual-marker annotation layer implemented in gloss.py. Together models "
     "sign grammar explicitly rather than treating translation as a bag-of-words "
     "substitution, which is what separates intelligible output from word salad.")

h2("7.2 Bidirectional Gloss Translation")
body("Two complementary transformations are needed. In the sign-to-text direction, "
     "a recognised gloss sequence in Topic–Comment order must become a fluent "
     "spoken-language sentence: STORE I GO → “I am going to the store.” In "
     "the text-to-sign direction, a spoken-language sentence must become an ordered "
     "gloss sequence: “What is your name?” → YOUR NAME WHAT. Both are "
     "implemented by few-shot prompting of an LLM through the provider abstraction, "
     "with a compact set of high-signal exemplars that teach the model the rules of "
     "the transformation — fronting the topic, dropping articles and copulas, using "
     "ME and YOU for first- and second-person objects, and appending wh-words.")
code_block([
    "# gloss → sentence (English)",
    "Gloss: STORE I GO       → Sentence: I am going to the store.",
    "Gloss: YOUR NAME WHAT   → Sentence: What is your name?",
    "Gloss: MOTHER LOVE ME   → Sentence: My mother loves me.",
    "",
    "# sentence → gloss (English)",
    "Sentence: Do you want to eat? → Gloss: YOU WANT EAT YOU",
])

h2("7.3 Graceful Degradation Without an LLM")
body("Because translation depends on an external language model, both directions are "
     "engineered to degrade gracefully. If no LLM is reachable, gloss-to-sentence "
     "simply returns the raw gloss joined by spaces, so the caller always receives "
     "something displayable. Sentence-to-gloss falls back to a deterministic "
     "rule-based glosser that uppercases tokens and strips English articles and "
     "copulas using a stop-word list; this does not reorder into full Topic–Comment "
     "form but keeps the text-to-sign pipeline functional offline. Arabic keeps "
     "every token in the fallback because a reliable offline stop-word list is not "
     "assumed. This design ensures that the loss of cloud connectivity downgrades "
     "quality rather than breaking the feature.")

h2("7.4 The Non-Manual-Marker Layer")
body("Section 2.2.4 explained that much sign grammar lives on the face and body. "
     "Together exposes this through a documented annotation hook. The function "
     "detect_sentence_type classifies a source sentence as a wh-question, a yes/no "
     "question, a negation, a topic, or a plain statement, using lexical cues "
     "(wh-words and negation words in both English and Arabic) and punctuation. The "
     "function annotate_nmm then attaches the appropriate marker bundle — for "
     "example, furrowed eyebrows and a forward head for a wh-question, or a headshake "
     "and a frown for negation — together with the span of gloss tokens over which "
     "the marker applies.")
code_block([
    "annotate_nmm([\"YOUR\", \"NAME\", \"WHAT\"], \"What is your name?\") = {",
    "  \"tokens\": [\"YOUR\", \"NAME\", \"WHAT\"],",
    "  \"sentence_type\": \"wh_question\",",
    "  \"markers\": {\"eyebrows\": \"furrowed\", \"head\": \"forward\"},",
    "  \"span\": [0, 3]   # marker applies across the whole clause",
    "}",
])
body("The marker vocabulary is centralised, and the sentence-type heuristics are "
     "deliberately isolated behind a single function so that a learned classifier "
     "can later replace the rules without changing any consumer. The avatar renderer "
     "is the intended consumer: it is expected to read the marker bundle and apply "
     "eyebrow, head, and mouth movements across the indicated span. This layer is a "
     "first-class architectural hook for facial grammar, even though full facial "
     "rendering on the avatar remains future work.")
table(
    ["Sentence type", "Cue", "Non-manual markers"],
    [
        ["wh_question", "what/where/when/who/why/how + Arabic equivalents", "eyebrows furrowed, head forward"],
        ["yes_no_question", "ends with “?”", "eyebrows raised, head forward, hold"],
        ["negation", "not/no/never + Arabic equivalents", "head shake, mouth frown"],
        ["topic", "fronted topic constituent", "eyebrows raised, head tilt"],
        ["statement", "default", "(none)"],
    ],
    caption="Sentence types, their lexical cues, and the non-manual markers attached "
            "by the annotation layer.")

h2("7.5 Caching Repeated Translations")
body("In conversation the same gloss recurs constantly, and an LLM round-trip is the "
     "dominant latency in the sign-to-text path. Gloss-to-sentence therefore caches "
     "its results in a bounded 256-entry cache keyed on the (gloss, language) pair, "
     "evicting the oldest entry when full. A repeated gloss is answered from cache in "
     "microseconds instead of incurring a fresh cloud round-trip of several hundred "
     "milliseconds. The cache is bypassed when a provider is injected explicitly (as "
     "in tests) to keep behaviour deterministic.")

# ═══════════════════════════ CHAPTER 8 ═══════════════════════════
h1("Chapter 8: Text and Speech to Sign")

h2("8.1 Introduction")
body("This chapter describes the reverse pipeline: turning a spoken-language "
     "sentence into a sequence of sign animations played by an avatar. The pipeline "
     "has four stages — optional transcription of speech to text, conversion of text "
     "to gloss with non-manual markers, mapping of each gloss token to a stored "
     "landmark sequence, and synthesis of any out-of-vocabulary words by "
     "fingerspelling and stitching.")

h2("8.2 Pipeline Overview")
body("When the input is speech, it is first transcribed by the speech-to-text "
     "provider chain. The resulting text (or directly typed text) is converted to "
     "ordered gloss tokens by the sentence-to-gloss translator of Chapter 7, "
     "annotated with non-manual markers. The client then requests a batch lookup of "
     "landmark sequences for the gloss tokens; for each token the server attempts an "
     "exact match, then a contraction/alias match, then a semantic match, and "
     "finally a stitched fingerspelling synthesis. The avatar plays the returned "
     "sequences in order, transitioning smoothly between them.")
figure_placeholder("End-to-end speech/text-to-sign pipeline: transcription → gloss "
                   "→ batched sign lookup → avatar playback.")

h2("8.3 Sign Lookup: Exact, Alias and Semantic Matching")
body("Mapping a gloss token to a stored sign is more subtle than a dictionary "
     "lookup because users and the gloss translator do not always produce the exact "
     "lemma stored in the database. The SignDB.match_word method applies a cascade. "
     "First it tries an exact match against the vocabulary, including a "
     "punctuation-stripped variant. Certain tokens that cause bad semantic matches — "
     "such as “i”, “am”, “is”, “be” — are "
     "explicitly skipped, and a small alias map redirects common contractions (for "
     "example “im” → “i am”). If no exact or alias match exists, "
     "the system falls back to semantic search.")
h3("8.3.1 Sentence-BERT Semantic Search over pgvector")
body("Semantic matching encodes the query word with Sentence-BERT (all-MiniLM-L6-v2) "
     "into a 384-dimensional embedding and finds the nearest stored sign by cosine "
     "distance. Crucially, this search is performed as an indexed pgvector query in "
     "PostgreSQL, using the HNSW index on the embedding column, rather than an "
     "in-Python brute-force scan over all embeddings. This keeps lookup fast and "
     "memory-light as the vocabulary grows, and means no embedding matrix is held in "
     "RAM. A maximum cosine-distance threshold (0.35 for English, 0.55 for the "
     "smaller, more permissive Arabic vocabulary) prevents spurious matches: if "
     "nothing falls within the threshold, the word is treated as out of vocabulary.")
code_block([
    "q_emb = sbert.encode(query)              # 384-d embedding",
    "hit = repo.nearest(q_emb, language, MAX_DISTANCE)  # indexed pgvector kNN",
    "return hit.word if hit else None         # else: out-of-vocabulary",
])

h2("8.4 Gloss-and-Stitch: Synthesising Out-of-Vocabulary Words")
body("No finite vocabulary can cover every word a user might say. For words with no "
     "dedicated sign — proper nouns especially — Together synthesises an animation by "
     "fingerspelling: it looks up the sign for each letter and concatenates the "
     "clips. The difficulty is that naïvely concatenating clips produces a visible "
     "jerk at every boundary, because the last pose of one letter rarely matches the "
     "first pose of the next. The stitch.py module solves this in two steps. First, "
     "it inserts a short linearly-interpolated bridge of a few frames between "
     "consecutive clips, so the hands glide from one letter's final pose to the "
     "next letter's initial pose. Second, it applies a Savitzky–Golay filter over a "
     "window centred on each stitch boundary, smoothing the transition while "
     "preserving the body of each sign.")
h3("8.4.1 The Savitzky–Golay Smoothing")
body("A Savitzky–Golay filter fits a low-order polynomial to a sliding window of "
     "samples by least squares and replaces the centre value with the fitted value, "
     "which smooths noise while preserving the shape of features better than a plain "
     "moving average. Together applies it only in a local window around each "
     "boundary, so the interior of each letter sign is left untouched, with a window "
     "length of seven frames and a polynomial order of two by default. A vital "
     "detail is the handling of missing landmarks: columns containing NaN (from "
     "dropped MediaPipe tracking) are left unchanged, because smoothing across a NaN "
     "would poison otherwise good neighbouring frames. If any required letter sign is "
     "missing from the database, the synthesiser returns nothing and the word is "
     "honestly reported as unrenderable rather than shown as a broken animation.")
code_block([
    "TRANSITION_FRAMES = 5    # interpolated bridge length at each boundary",
    "SMOOTH_WINDOW     = 7    # Savitzky–Golay window (odd)",
    "SMOOTH_POLYORDER  = 2    # polynomial order",
    "# 1) insert interpolated bridge between clip A and clip B",
    "# 2) Savitzky–Golay smooth a window around each boundary (skip NaN columns)",
])
figure_placeholder("Stitching two fingerspelled letter clips: an interpolated bridge "
                   "plus Savitzky–Golay smoothing at the boundary removes the jerk.")

h2("8.4b Batched Lookup and Order Preservation")
body("A naturally signed sentence contains several tokens, and resolving each one "
     "independently with a separate request would be wasteful and would risk "
     "delivering the signs out of order. Together therefore exposes a batch endpoint "
     "that accepts the whole ordered gloss sequence and returns, for each token, the "
     "resolved landmark sequence together with the source of the resolution — exact, "
     "alias, semantic, or stitched. Preserving order is essential: the avatar must "
     "sign “MOTHER LOVE ME” in that order, and a Topic–Comment construction "
     "loses its meaning if the topic is not signed first. The batch response is "
     "therefore index-aligned with the input gloss, and the client concatenates the "
     "sequences in order, inserting the same kind of smooth transition between whole "
     "signs that the stitcher inserts between fingerspelled letters. Tagging each "
     "result with its source lets the interface communicate uncertainty honestly — "
     "for example, indicating when a word was fingerspelled because no lexical sign "
     "existed, or when a semantic match was used rather than an exact one.")
body("This batching also interacts well with the performance work of Chapter 13: "
     "the entire lookup runs off the event loop in a thread pool, the SBERT encoding "
     "of any semantically matched tokens happens together, and the indexed pgvector "
     "queries are individually fast, so a multi-word sentence resolves with low total "
     "latency rather than paying a separate round-trip per word.")

h2("8.5 The Avatar and Landmark Store")
body("Landmark sequences are stored as compressed .npz files on disk and loaded on "
     "demand by the landmark store, keyed by sign id, so that only the signs a "
     "request actually needs are read into memory. The client receives each sign's "
     "[n_frames, 543, 3] sequence and drives the avatar to reproduce the motion, "
     "concatenating the per-token sequences (and any stitched synthesis) into one "
     "continuous signed utterance. Stitched results are tagged with their source so "
     "the interface can indicate when a word was fingerspelled rather than signed "
     "lexically.")

# ═══════════════════════════ CHAPTER 9 ═══════════════════════════
h1("Chapter 9: The Provider Abstraction")

h2("9.1 Introduction")
body("Together depends on three external capabilities — language modelling for gloss "
     "translation, text-to-speech for spoken output, and speech-to-text for spoken "
     "input. Hard-wiring a single vendor would make the system fragile (a network "
     "outage breaks it), expensive (no free path), and hard to evolve. This chapter "
     "describes the pluggable provider abstraction that isolates the rest of the "
     "application from any concrete vendor and that degrades gracefully from cloud "
     "services to fully offline fallbacks.")

h2("9.2 Abstract Interfaces")
body("The provider layer defines three abstract base classes. LLMProvider exposes a "
     "single generate(prompt, temperature) method; TTSProvider exposes "
     "synthesize(text, language) returning WAV bytes; and STTProvider exposes "
     "transcribe(audio, mime_type, language). Each also offers a cheap availability "
     "check. The rest of the application — the gloss engine, the speech endpoints — "
     "depends only on these interfaces and never on a concrete vendor. Concrete "
     "adapters for Gemini, Ollama, and the local engines implement the interfaces, "
     "and a factory wires the chain together based on environment variables. A common "
     "ProviderError signals that an adapter could not fulfil a request, which the "
     "fallback chain catches to try the next provider in line.")
code_block([
    "class LLMProvider(ABC):",
    "    def generate(self, prompt, temperature=0.0) -> str: ...",
    "class TTSProvider(ABC):",
    "    def synthesize(self, text, language='english') -> bytes: ...",
    "class STTProvider(ABC):",
    "    def transcribe(self, audio, mime_type='audio/wav', language='english') -> str: ...",
])

h2("9.3 The Fallback Chains")
body("Each capability is configured with a primary provider and an automatic offline "
     "fallback. The defaults are summarised below. When no cloud key is set, or when "
     "a cloud call raises a ProviderError, the next link in the chain takes over so "
     "the user-facing feature keeps working — at possibly lower quality — rather than "
     "failing outright. Each fallback can be individually disabled to force hard "
     "failures in production environments that prefer explicit errors over silent "
     "degradation.")
table(
    ["Capability", "Primary (cloud)", "Offline fallback (free, local)"],
    [
        ["LLM (gloss ↔ sentence)", "Google Gemini", "Ollama (e.g. llama3.2)"],
        ["TTS (speech output)", "Google Gemini TTS", "pyttsx3 (OS voices)"],
        ["STT (speech input)", "Google Gemini", "faster-whisper (local)"],
    ],
    caption="Provider chains: each capability prefers a cloud provider and degrades "
            "to a free, local fallback.")

h2("9.4 Cost and Operational Considerations")
body("The only paid API in the default configuration is Google Gemini, and its "
     "free tier on the flash models is sufficient for development and light demo "
     "usage. Moving to production simply means swapping in a paid Gemini key — no "
     "code changes — or, for an air-gapped deployment, configuring the offline "
     "providers exclusively. The offline fallbacks (Ollama, pyttsx3, Whisper) are "
     "free and run locally, trading cloud cost and dependency for local compute. "
     "This design lets the same codebase serve a zero-cost classroom demo, a "
     "cloud-backed production service, and a fully offline air-gapped deployment "
     "merely by changing environment variables.")
table(
    ["Scenario", "Configuration"],
    [
        ["Development / demo", "Gemini free-tier key; fallbacks enabled"],
        ["Production, cloud", "Paid Gemini key; fallbacks optionally disabled for hard failures"],
        ["Fully offline / air-gapped", "Ollama + pyttsx3 + Whisper; no Gemini key"],
    ],
    caption="Recommended provider configurations by deployment scenario.")

# ═══════════════════════════ CHAPTER 10 ═══════════════════════════
h1("Chapter 10: Real-Time Meeting Mode")

h2("10.1 Introduction")
body("The translation modules described so far serve a single user at a time. The "
     "meeting mode brings two people together: a signer and a speaker, each in their "
     "own browser, conversing live with translation running continuously for both. "
     "This chapter describes the real-time architecture that makes this possible — "
     "WebRTC for peer-to-peer media and Socket.IO for signalling — and the "
     "role-based interface that presents the right view to each participant.")

h2("10.2 WebRTC and Peer-to-Peer Media")
body("WebRTC (Web Real-Time Communication) is the browser standard for "
     "low-latency, peer-to-peer audio and video. Crucially, once a connection is "
     "established, media flows directly between the two browsers rather than through "
     "the application server, which minimises latency and server bandwidth and keeps "
     "the conversation private. Establishing the connection, however, requires the "
     "two peers to exchange session descriptions (offers and answers) and network "
     "candidates (ICE), and they cannot do so directly before the connection exists. "
     "A signalling channel is therefore needed to broker this initial handshake.")
figure_placeholder("WebRTC connection establishment: offer/answer and ICE candidate "
                   "exchange via the Socket.IO signalling server, after which media "
                   "flows peer-to-peer.")

h2("10.3 Socket.IO Signalling")
body("Together uses a Socket.IO server, mounted alongside the FastAPI application, "
     "as the signalling channel. The server handles a small set of events: joining "
     "and leaving a room, announcing presence so peers discover one another, and "
     "relaying the WebRTC offer, answer, and ICE candidates between the two members "
     "of a room. It also relays translated sentences for the live-caption path. "
     "Because only these small control messages traverse the server, the signalling "
     "load is negligible compared with media, and the server never sees the "
     "conversational audio or video.")
table(
    ["Socket.IO event", "Purpose"],
    [
        ["join_room / leave_room", "Manage room membership"],
        ["announce_presence / *_reply", "Peer discovery within a room"],
        ["webrtc_offer", "Relay an SDP offer from caller to callee"],
        ["webrtc_answer", "Relay an SDP answer from callee to caller"],
        ["webrtc_ice_candidate", "Relay ICE network candidates between peers"],
        ["translate_sentence", "Relay a translated caption/sentence to the room"],
    ],
    caption="Socket.IO signalling events used by the meeting mode.")

h2("10.4 Role-Based Interfaces")
body("A meeting has two asymmetric roles. The Signer signs to the camera and wants "
     "to read or see what the Speaker says; the Speaker talks and wants to "
     "understand what the Signer signs. The interface adapts to the role: the Signer "
     "sees captions of the Speaker's transcribed and translated speech (or the "
     "avatar), while the Speaker sees captions of the Signer's recognised, "
     "gloss-translated signing. The user's role is part of their profile, and the "
     "translation direction for each participant follows from it, so the same "
     "meeting runs both pipelines simultaneously, one per direction.")
figure_placeholder("Role-based meeting interface: the Signer's and Speaker's "
                   "complementary views.")

h2("10.4b NAT Traversal: STUN and TURN")
body("A practical complication of peer-to-peer media is that most devices sit behind "
     "Network Address Translation (NAT) and firewalls, so neither peer knows a "
     "routable address at which the other can be reached. WebRTC solves this through "
     "the Interactive Connectivity Establishment (ICE) framework, which gathers "
     "candidate addresses and tests them for connectivity. A STUN (Session Traversal "
     "Utilities for NAT) server lets a peer discover its own public-facing address "
     "and port as seen from the outside, which is sufficient for the common case "
     "where at least one side's NAT is cooperative. When both peers are behind "
     "restrictive symmetric NATs and no direct path can be found, a TURN (Traversal "
     "Using Relays around NAT) server relays the media as a last resort. Together's "
     "signalling exchanges the ICE candidates these mechanisms produce; the candidate "
     "exchange seen by the Socket.IO layer is precisely the output of this discovery "
     "process. In deployment, configuring reliable STUN and, for difficult networks, "
     "TURN servers is essential for connection success rates, and is a known "
     "operational cost of any WebRTC application.")

h2("10.5 Resilience")
body("Real-time connections are inherently fragile — networks change, NATs "
     "intervene, and peers disconnect. The signalling layer is structured so that "
     "presence can be re-announced and the offer/answer exchange repeated to "
     "re-establish a dropped connection. Because the translation pipelines are "
     "stateless per utterance, a brief media interruption does not corrupt the "
     "conversation; recognition simply resumes when frames flow again.")

# ═══════════════════════════ CHAPTER 11 ═══════════════════════════
h1("Chapter 11: Authentication, Security and Data Layer")

h2("11.1 Introduction")
body("A system that handles user accounts, camera and microphone access, and live "
     "conversations must be secure by design. This chapter describes Together's "
     "authentication and authorisation mechanisms — Argon2id password hashing, JWT "
     "access and refresh tokens, and rate limiting — and the PostgreSQL/pgvector "
     "data layer that backs both the user accounts and the semantic sign search.")

h2("11.2 Password Hashing with Argon2id")
body("Passwords are never stored in plaintext. Together hashes them with Argon2id, "
     "the winner of the Password Hashing Competition and the current OWASP "
     "recommendation, configured with a time cost of three, a memory cost of 64 MB, "
     "and a parallelism of two — parameters chosen to make brute-force attacks "
     "expensive while keeping legitimate logins fast. The implementation also retains "
     "a bcrypt verifier purely to validate any legacy hashes during a transparent "
     "migration: a needs_rehash check detects an old-scheme hash on successful login "
     "and the password is silently upgraded to Argon2id, so the user base migrates "
     "without anyone resetting a password.")
code_block([
    "PasswordHasher(time_cost=3, memory_cost=65536, parallelism=2,",
    "               hash_len=32, salt_len=16)   # Argon2id, OWASP params",
    "# legacy bcrypt hashes ($2b$/$2a$/$2y$) verified, then rehashed to Argon2id",
])

h2("11.3 JSON Web Tokens: Access and Refresh")
body("Sessions are managed with two JSON Web Tokens. A short-lived access token "
     "(default sixty minutes) authorises API requests and is sent as a bearer token. "
     "A long-lived refresh token (default seven days) is used only to obtain new "
     "access tokens. Critically, refresh tokens are not merely signed and trusted: "
     "each is registered server-side in a refresh-token table by a unique identifier "
     "(jti), so it can be revoked. On every refresh the old token is validated, "
     "revoked, and a brand-new token is issued — a rotation scheme that limits the "
     "damage of a stolen refresh token, because using an already-rotated token is "
     "rejected. Logout revokes the token. Access tokens are explicitly distinguished "
     "from refresh tokens by a type claim, so a refresh token cannot be used to "
     "authorise an ordinary API call.")
figure_placeholder("JWT access/refresh token lifecycle with server-side rotation "
                   "and revocation.")
table(
    ["Mechanism", "Purpose", "Detail"],
    [
        ["Access token", "Authorise API calls", "HS256-signed JWT, ~60 min, bearer header, type=access"],
        ["Refresh token", "Obtain new access tokens", "~7 days, registered by jti, revocable, type=refresh"],
        ["Rotation", "Limit stolen-token damage", "Old token revoked and replaced on each refresh"],
        ["Logout", "End a session", "Refresh token revoked server-side"],
    ],
    caption="Token mechanisms and their security properties.")

h2("11.4 Rate Limiting")
body("To blunt brute-force and credential-stuffing attacks, the authentication "
     "endpoints are protected by an in-memory sliding-window rate limiter that "
     "permits at most ten requests per sixty seconds per client IP, returning HTTP "
     "429 with a Retry-After header when exceeded. The sliding window is more "
     "accurate than a fixed bucket because it counts only the requests within the "
     "trailing window rather than resetting on a boundary.")

h2("11.5 The PostgreSQL and pgvector Data Layer")
body("The persistent store is PostgreSQL with the pgvector extension. Three ORM "
     "entities model the domain. User holds the account, the Argon2id hash, a role, "
     "and a creation timestamp. RefreshToken implements the rotation and revocation "
     "registry described above. Sign records one vocabulary item per language with "
     "its video filename, the on-disk landmark file path, a frame count, and a "
     "384-dimensional pgvector embedding. The Sign table enforces a unique constraint "
     "on the (word, language) pair and, decisively, builds an HNSW index on the "
     "embedding column with cosine-distance operators so that semantic search is a "
     "fast approximate nearest-neighbour query in the database itself.")
table(
    ["Entity", "Key fields", "Notes"],
    [
        ["User", "id, email (unique), hashed_password, role, created_at", "Argon2id hash; role drives meeting view"],
        ["RefreshToken", "jti (unique), user_id, revoked, expires_at", "Server-side rotation/revocation"],
        ["Sign", "word, language, video_filename, landmark_file, frame_count, embedding(384)", "Unique (word, language); HNSW cosine index"],
    ],
    caption="The Together data model.")
body("Storing the 384-dimensional embeddings in the database with an HNSW index, "
     "while keeping the bulky [n_frames, 543, 3] landmark arrays as compressed files "
     "on disk, is the design that lets semantic search scale without holding any "
     "embedding matrix in memory and without bloating the database with motion data.")

h2("11.5b Threat Model and Mitigations")
body("Security decisions are only meaningful relative to the threats they address. "
     "The table below states the principal threats Together considers and the "
     "specific mitigation each receives, so that the security design can be audited "
     "against a concrete adversary model rather than assessed in the abstract.")
table(
    ["Threat", "Mitigation in Together"],
    [
        ["Offline password cracking of a leaked database", "Argon2id memory-hard hashing (64 MB, t=3) makes brute force expensive"],
        ["Credential stuffing / brute-force login", "Sliding-window rate limit: 10 requests / 60 s / IP, HTTP 429"],
        ["Stolen refresh token replayed", "Server-side rotation: each refresh revokes the old jti; replay is rejected"],
        ["Refresh token used as an access token", "Explicit type claim distinguishes access from refresh tokens"],
        ["Token forgery", "HS256 signature over a server-held secret from the environment"],
        ["Cross-origin request abuse", "Configurable CORS allow-list"],
        ["Injection / malformed input", "Pydantic schema validation and email validation"],
        ["Information leakage in errors", "Generic and validation exception handlers hide internals"],
        ["Privacy exposure of camera feed", "Only derived landmarks leave the browser; raw video never does"],
        ["Secret leakage via source control", "Secrets supplied through environment variables, not committed"],
    ],
    caption="Threat model: principal threats and their mitigations.")

h2("11.6 Additional Hardening")
body("Beyond authentication, the platform applies standard web hardening: a "
     "configurable CORS allow-list restricts which origins may call the API; input "
     "is validated through Pydantic schemas and email validation; generic and "
     "validation exception handlers prevent internal details from leaking in error "
     "responses; and the JWT signing secret, database DSN, and provider keys are all "
     "supplied through environment variables rather than committed to source.")

# ═══════════════════════════ CHAPTER 12 ═══════════════════════════
h1("Chapter 12: Frontend and Design System")

h2("12.1 Introduction")
body("The most sophisticated translation engine is useless if people cannot operate "
     "it. Together's front end is a deliberately simple, server-rendered, "
     "framework-free interface built on a design-token system that makes the UI "
     "fully bilingual, themeable, and accessible. This chapter describes that design "
     "system, the bilingual and right-to-left support, and the accessibility "
     "features.")

h2("12.2 A Framework-Free, Server-Rendered Front End")
body("The interface is composed of Jinja2 templates rendering plain HTML, styled by "
     "a small CSS system and driven by vanilla JavaScript — no single-page-application "
     "framework and no build step. This choice keeps the client lightweight and "
     "fast to load, eliminates a large class of build-tooling complexity, and makes "
     "the markup easy to mirror for right-to-left languages. The landing, "
     "authentication, and dashboard pages share a common base template, while the "
     "English and Arabic dashboards share structure and inline logic but differ in "
     "direction, default model, and copy.")

h2("12.3 The Design-Token System")
body("Styling is centralised in a design-token system so that a restyle is a "
     "one-file change and light, dark, and RTL variants come essentially for free. A "
     "single theme stylesheet is the source of truth for colour, typography, "
     "spacing, and elevation tokens. Light and dark themes are simple token swaps "
     "keyed on a data attribute on the root element, with a no-JavaScript fallback "
     "that honours the operating system's preferred colour scheme. A separate "
     "component stylesheet defines reusable, prefixed UI components (buttons, cards, "
     "pills, sidebars, toggles), and a small no-flash theme controller applies the "
     "stored or system theme synchronously before the first paint to avoid a flash "
     "of the wrong theme. Brand fonts are bundled locally rather than fetched from a "
     "content-delivery network, removing an external dependency.")
figure_placeholder("The Together design system: light and dark themes derived from "
                   "the same token set.")

h2("12.4 Bilingual and Right-to-Left Support")
body("Arabic is a right-to-left language, and a half-hearted RTL implementation "
     "produces a broken, mirrored-but-wrong layout. Together handles RTL with CSS "
     "logical properties — inline-start and inline-end rather than left and right — "
     "so that the entire interface mirrors correctly at zero additional cost when "
     "the document direction is set to RTL. The Arabic dashboard is right-to-left, "
     "forces the Arabic recognition model, and bakes Arabic copy directly into the "
     "markup, while sharing its structure with the English dashboard. The result is "
     "two genuinely first-class language experiences from one shared codebase.")
figure_placeholder("Side-by-side English (LTR) and Arabic (RTL) dashboards "
                   "demonstrating mirrored layout.")

h2("12.5 Accessibility")
body("Because the audience includes people with disabilities, accessibility is a "
     "requirement rather than a nicety. The interface provides visible focus rings "
     "for keyboard navigation, a skip-link to bypass navigation, and honours the "
     "user's reduced-motion preference to avoid triggering vestibular discomfort. "
     "Semantic markup and adequate colour contrast in both themes support assistive "
     "technologies. These features, together with the privacy guarantee that no raw "
     "video leaves the device, make the system respectful of the people it is "
     "designed to serve.")

h2("12.6 The Translation Dashboards")
body("Each dashboard hosts the four translation modules behind a single, coherent "
     "interface: a camera view with the avatar, controls to select the module "
     "(sign-to-text, sign-to-speech, text-to-sign, speech-to-sign) and the language, "
     "and a transcript area. The dashboards load the MediaPipe and WebAssembly "
     "runtime, connect the Socket.IO client for meetings, and wire the camera and "
     "microphone to the appropriate endpoints. The interface degrades gracefully: if "
     "a capability's provider is unavailable, the corresponding output downgrades "
     "rather than disappearing.")
figure_placeholder("The main translation dashboard with camera, avatar, module "
                   "controls, and transcript.")

# ═══════════════════════════ CHAPTER 13 ═══════════════════════════
h1("Chapter 13: Performance Engineering and Latency")

h2("13.1 Introduction")
body("A conversational system lives or dies by its latency. A translation that "
     "arrives two seconds late breaks the rhythm of a conversation. This chapter "
     "documents the profiling infrastructure that makes latency measurable and the "
     "set of optimisations that keep Together responsive under load. The guiding "
     "principle is that one cannot improve what one cannot measure, so observability "
     "comes first.")

h2("13.2 Profiling Infrastructure")
body("Every hot stage of the pipeline is wrapped in a lightweight timer, and the "
     "measurements are aggregated and exposed at a metrics endpoint that reports, per "
     "stage, the count and the average, 50th-percentile, 95th-percentile, maximum, "
     "and last latency in milliseconds. Stages tracked include ASL and Arabic "
     "inference, the gloss-to-sentence and sentence-to-gloss LLM calls, speech "
     "transcription, and the sign-lookup operations. This makes performance "
     "regressions visible without any external tooling and turns optimisation from "
     "guesswork into a measured activity. Offline micro-benchmarks complement the "
     "live metrics for components that can be exercised without a camera or network.")
table(
    ["Stage key", "What it measures"],
    [
        ["infer.asl", "TensorFlow-Lite ASL inference"],
        ["infer.arabic", "PyTorch CNN–GRU Arabic inference"],
        ["llm.gloss_to_sentence", "LLM gloss → sentence translation"],
        ["llm.english_to_gloss", "LLM sentence → gloss translation"],
        ["stt.transcribe", "Speech-to-text transcription"],
        ["signs.*", "Sign lookup and batch retrieval"],
    ],
    caption="Per-stage timings exposed by the metrics endpoint "
            "(count / avg / p50 / p95 / max / last).")

h2("13.3 Optimisation 1: Moving Blocking Work Off the Event Loop")
body("The most important optimisation is architectural. The asynchronous endpoints "
     "originally ran blocking work — TensorFlow-Lite and PyTorch inference, network "
     "LLM and STT calls, and SBERT encoding — directly in the event loop, so a "
     "single in-flight inference serialised every other request, including health "
     "checks and the meeting socket. All such blocking calls were moved onto a thread "
     "pool via Starlette's run_in_threadpool. The effect is that under concurrent "
     "load the event loop stays responsive and the 95th-percentile latency of "
     "unrelated requests no longer tracks the slowest inference.")

h2("13.4 Optimisation 2: INT8 Quantization of the Arabic Model")
body("As introduced in Chapter 6, the Arabic model's Linear and GRU weights are "
     "dynamically quantised to INT8 on CPU, and the inference path uses PyTorch's "
     "inference_mode to drop autograd bookkeeping. Measured on CPU with synthetic "
     "input over fifty runs, this reduces per-inference time from 7.48 ms (FP32 with "
     "inference_mode) to 4.21 ms (INT8 dynamic), a 1.78× speed-up, with negligible "
     "accuracy impact for a model of this size.")
table(
    ["Configuration", "Latency per inference", "Speed-up"],
    [
        ["FP32 (inference_mode)", "7.48 ms", "1.00×"],
        ["INT8 dynamic quantization", "4.21 ms", "1.78×"],
    ],
    caption="Measured Arabic-model inference latency on CPU (50 runs, synthetic "
            "[60, 543, 3] input).")

h2("13.5 Optimisation 3: TensorFlow-Lite Threading and XNNPACK")
body("The LiteRT interpreter for the ASL model is constructed with a thread count of "
     "the minimum of four and the available cores, overridable by environment "
     "variable, and LiteRT applies the XNNPACK delegate by default for float CPU "
     "models, accelerating the convolution and dense operations that dominate "
     "inference. The implementation falls back gracefully on older interpreter "
     "signatures that do not accept a thread-count argument.")

h2("13.6 Optimisation 4: Caching Repeated Gloss Translations")
body("Because the same gloss recurs constantly in conversation and the LLM "
     "round-trip is the dominant cost in the sign-to-text path, the gloss-to-sentence "
     "result is cached in a bounded cache keyed on (gloss, language). The measured "
     "effect is dramatic: a cold call with a simulated 50 ms LLM costs about 50 ms, "
     "whereas a warm cache hit costs about one microsecond, effectively eliminating "
     "the cost of repeated translations and, in production, removing a full cloud "
     "round-trip of several hundred milliseconds for any repeated gloss.")
table(
    ["Call", "Latency"],
    [
        ["Cold (50 ms simulated LLM)", "50.13 ms"],
        ["Warm (cache hit)", "0.001 ms"],
    ],
    caption="Gloss-to-sentence cache: cold versus warm latency.")

h2("13.7 Optimisation 5: Bounded Sliding Window")
body("Inference always operates on a bounded rolling window of frames — the ASL path "
     "trims to 60 frames and the Arabic path samples fixed 30/45/60-frame windows — "
     "so the cost of inference does not grow as a user keeps signing and the payload "
     "lengthens. This bounds the per-inference cost regardless of session length.")

h2("13.7b Latency Budget of an End-to-End Translation")
body("It is instructive to assemble the individual stage costs into an end-to-end "
     "budget for a single sign-to-speech translation, because this reveals where the "
     "user-perceived delay actually accumulates and validates the prioritisation of "
     "the optimisations above. Landmark extraction runs continuously in the browser "
     "and overlaps with everything else, so it does not add to the critical path of a "
     "completed utterance. The recognition inference is a few milliseconds on CPU "
     "after the optimisations, and the debouncing buffer adds a deliberate, "
     "human-scale delay equal to roughly the time a sign is held. The dominant "
     "variable cost is the LLM gloss-to-sentence call: several hundred milliseconds "
     "for a cold cloud round-trip, but effectively zero for a cached repeat and "
     "non-blocking for fresh calls. Speech synthesis adds the final stage for "
     "sign-to-speech. The clear conclusion, consistent with the measurements, is "
     "that the network language call dominates the machine cost, which is exactly why "
     "it is both cached and offloaded from the event loop.")
table(
    ["Stage", "Typical cost", "On critical path?"],
    [
        ["Landmark extraction (browser)", "Continuous, overlapped", "No"],
        ["Recognition inference (CPU)", "~4–8 ms", "Yes (small)"],
        ["Debouncing buffer", "~ sign-hold duration", "Yes (intentional)"],
        ["Gloss → sentence (LLM)", "~300–800 ms cold; ~0 cached", "Yes (dominant)"],
        ["Speech synthesis (TTS)", "Sub-second", "Yes (sign→speech only)"],
    ],
    caption="Approximate end-to-end latency budget for a sign-to-speech translation.")

h2("13.8 Net Effect")
body("Taken together, these optimisations produce a system that stays responsive "
     "under concurrency because no endpoint blocks the event loop; that runs Arabic "
     "recognition roughly 1.8× faster on CPU; that answers repeated translations "
     "almost instantly from cache; and whose remaining costs are visible at a glance "
     "through the metrics endpoint, so future regressions are caught early. The "
     "dominant residual cost in the sign-to-text path is the cold LLM call, which is "
     "both cached for repeats and offloaded from the event loop for fresh calls.")

# ═══════════════════════════ CHAPTER 14 ═══════════════════════════
h1("Chapter 14: Evaluation and Testing")

h2("14.1 Introduction")
body("This chapter evaluates Together along two axes: the recognition accuracy of "
     "the models, assessed against a held-out set of sign videos, and the "
     "correctness and robustness of the software, assessed through an automated test "
     "suite. It also reflects critically on the qualitative behaviour of the "
     "end-to-end system.")

h2("14.2 Model Evaluation Methodology")
body("The ASL recognition model was evaluated against a corpus of sign videos using "
     "an offline harness that runs each video through the same landmark-extraction "
     "and inference path as the live system, aligning ground-truth labels to the "
     "exact label map used by the deployed model so that the measurement reflects "
     "real behaviour. Of 272 video files, 250 mapped cleanly to model labels and "
     "were evaluated; 22 were skipped as unmapped, and none failed processing. The "
     "primary metric is Top-1 accuracy — the fraction of samples whose top "
     "prediction equals the ground truth — supplemented by average softmax "
     "confidence and the average top logit.")

h2("14.3 Model Baseline Results")
table(
    ["Metric", "Value"],
    [
        ["Total video files", "272"],
        ["Mapped / evaluated samples", "250"],
        ["Skipped (unmapped)", "22"],
        ["Failed processing", "0"],
        ["Top-1 accuracy", "0.6240 (62.4%)"],
        ["Average softmax confidence", "0.5450"],
        ["Average top logit", "6.4758"],
        ["Acceptance rate", "1.0000"],
        ["Accepted precision", "0.6240"],
        ["Runtime", "1021.4 s (4.09 s/sample)"],
    ],
    caption="ASL model baseline evaluation over 250 mapped samples.")
body("A Top-1 accuracy of 62.4% over a 250-class vocabulary is a strong result for "
     "isolated-sign recognition from landmarks alone: random guessing over 250 "
     "classes would score 0.4%, so the model performs roughly 156× better than "
     "chance. The acceptance rate of 1.0 with accepted precision equal to overall "
     "accuracy indicates that, at the evaluation threshold, the model commits to a "
     "prediction on every sample. The average softmax confidence of 0.545 reflects "
     "the genuine ambiguity among visually similar signs in a large vocabulary.")
figure_placeholder("Top-1 accuracy of the ASL model relative to chance over the "
                   "250-class vocabulary.")
h3("14.3.1 Confusion Analysis")
body("Examining the confusion candidates is instructive: the errors are dominated by "
     "single-count confusions between signs that are genuinely similar in handshape "
     "or motion — for example after/arm, all/weus, alligator/clean, because/for, and "
     "before/have. The long tail of one-off confusions, rather than a few "
     "systematically failing classes, suggests that the principal route to higher "
     "accuracy is more and more varied training data and finer temporal modelling, "
     "rather than a structural defect in the model.")
table(
    ["Ground truth", "Predicted", "Count"],
    [
        ["after", "arm", "1"],
        ["all", "weus", "1"],
        ["alligator", "clean", "1"],
        ["because", "for", "1"],
        ["before", "have", "1"],
        ["better", "arm", "1"],
        ["boat", "book", "1"],
    ],
    caption="Representative top confusion candidates, dominated by visually similar "
            "sign pairs.")

h2("14.4 Arabic Model Performance")
body("The Arabic CNN–GRU model recognises a curated 20-word vocabulary (baby, eat, "
     "father, finish, good, happy, hear, house, important, love, mall, me, mosque, "
     "mother, normal, sad, stop, thanks, thinking, worry). Its multi-window ensemble "
     "and confidence threshold of 0.65 favour precision over recall, emitting a "
     "prediction only when the averaged probability is decisive. As reported in "
     "Chapter 13, INT8 quantization makes it roughly 1.8× faster than its FP32 "
     "baseline on CPU while preserving accuracy, which is what makes real-time "
     "Arabic recognition feasible on commodity hardware.")

h2("14.5 Software Testing")
body("Correctness is guarded by an automated pytest suite spanning the platform. The "
     "authentication tests cover Argon2id hashing, the legacy-bcrypt verification and "
     "rehash path, JWT creation and validation, and the sliding-window rate limiter. "
     "The gloss tests cover bidirectional translation, the rule-based offline "
     "fallback, sentence-type detection, and non-manual-marker annotation. The "
     "provider tests verify the fallback chains. The stitch tests verify that "
     "concatenation, interpolation, and Savitzky–Golay smoothing produce well-formed "
     "sequences and that missing letters yield an honest None. Profiling and "
     "sentence-stitching paths are exercised as well.")
h3("14.5.1 Frontend Regression Tests")
body("A dedicated frontend regression module renders every template and asserts the "
     "invariants that had broken during a major redesign: that the dashboards retain "
     "their WebAssembly, MediaPipe, and Socket.IO head blocks; that the design-system "
     "stylesheets remain linked; that the legacy JavaScript hooks still exist; that "
     "every element looked up by id actually resolves in the markup; that the Arabic "
     "dashboard is right-to-left with no English leakage; and that the inline "
     "JavaScript parses. These tests turn a class of silent UI breakages into loud, "
     "automated failures.")
table(
    ["Test module", "Coverage"],
    [
        ["test_auth", "Hashing, tokens, rate limiting"],
        ["test_gloss", "Gloss translation, fallback, NMM"],
        ["test_providers", "Provider fallback chains"],
        ["test_stitch", "Interpolation and smoothing"],
        ["test_templates", "Frontend regression invariants"],
        ["test_profiling", "Timing instrumentation"],
    ],
    caption="The automated test suite by module.")

h2("14.6 Qualitative and Usability Evaluation")
body("Beyond quantitative accuracy, the system was exercised qualitatively to assess "
     "whether the end-to-end experience is genuinely usable. In informal sessions, "
     "the sign-to-speech flow was tested by signing short, in-vocabulary utterances "
     "and confirming that the debouncing buffer produced one clean gloss per held "
     "sign and that the assembled sentence was natural. The text-to-sign flow was "
     "tested by typing and speaking sentences and confirming that in-vocabulary words "
     "were signed lexically, that out-of-vocabulary proper nouns were fingerspelled "
     "smoothly without visible jerks at letter boundaries, and that genuinely "
     "unrenderable words were reported honestly rather than shown broken. The meeting "
     "mode was tested across two browsers to confirm peer-to-peer media establishment "
     "and per-role captioning.")
body("Several qualitative observations emerged. Recognition is most reliable under "
     "good, even lighting with the signer centred and the hands well separated from "
     "the body, consistent with the landmark-extraction front end. The confidence "
     "thresholds successfully suppress spurious output during transitions, at the "
     "cost of occasionally requiring a sign to be held a moment longer. The "
     "gloss-to-sentence LLM stage is forgiving of minor recognition noise, often "
     "assembling a sensible sentence even when one gloss is marginal. The offline "
     "fallbacks were validated by removing the cloud key and confirming that "
     "translation, synthesis, and transcription continued at reduced quality rather "
     "than failing. These findings corroborate the design intent: the system favours "
     "honest, conversational behaviour over brittle perfection.")
table(
    ["Flow", "Observed behaviour", "Verdict"],
    [
        ["Sign → speech (in-vocab)", "One clean gloss per held sign; natural sentence", "Works well"],
        ["Text → sign (in-vocab)", "Lexical signing of known words", "Works well"],
        ["Text → sign (OOV)", "Smooth fingerspelling; honest unrenderable report", "Works as designed"],
        ["Meeting mode", "P2P media established; role-based captions", "Works"],
        ["Offline fallback", "Reduced-quality continuation, no failure", "Resilient"],
    ],
    caption="Summary of qualitative end-to-end observations.")

h2("14.7 Critical Discussion")
body("The evaluation confirms that the system meets its functional and "
     "non-functional goals: it translates in both directions for two languages, it "
     "stays responsive under load, it degrades gracefully without the cloud, and it "
     "is secured and tested to a production standard. The honest limitation is "
     "recognition accuracy on the large ASL vocabulary, which at 62.4% Top-1 is good "
     "for the approach but not yet at the level required for unsupervised, "
     "high-stakes use; the debouncing buffer and confidence gating mitigate this in "
     "practice by suppressing low-confidence emissions, and the gloss-to-sentence LLM "
     "stage can absorb minor recognition noise when assembling a sentence. These "
     "observations set the agenda for the future work of the next chapter.")

# ═══════════════════════════ CHAPTER 15 ═══════════════════════════
h1("Chapter 15: Conclusion and Future Work")

h2("15.1 Conclusion")
body("This thesis set out to answer how one might build an accessible, low-cost, "
     "real-time system for two-way communication between sign-language users and "
     "non-signers, in both American and Arabic sign languages, using only a browser "
     "and commodity hardware, while respecting the grammar of sign languages and "
     "remaining usable without a constant cloud connection. Together answers that "
     "question concretely and completely.")
body("The system delivers all four translation flows — sign-to-text, sign-to-speech, "
     "text-to-sign, and speech-to-sign — for two first-class languages from a single "
     "install-free web application. It recognises signs from privacy-preserving "
     "MediaPipe landmarks using a 250-class TensorFlow-Lite ASL model and a "
     "purpose-built CNN–GRU Arabic model, debounced into clean glosses. It respects "
     "sign grammar through bidirectional, Topic–Comment-aware gloss translation and a "
     "documented non-manual-marker layer. It synthesises sign animations through "
     "indexed Sentence-BERT semantic lookup and a Savitzky–Golay-smoothed "
     "fingerspelling stitcher for out-of-vocabulary words. It connects two people "
     "live over peer-to-peer WebRTC with Socket.IO signalling and role-based views. "
     "And it does all of this resiliently, degrading from cloud providers to fully "
     "offline fallbacks so the core experience never breaks, on a secure, observable, "
     "and tested platform.")
body("The measured results substantiate the design: 62.4% Top-1 accuracy over a "
     "250-class ASL vocabulary — roughly 156× better than chance — a 1.78× speed-up "
     "from INT8 quantization of the Arabic model, and the near-total elimination of "
     "repeated-translation latency through caching. Beyond the numbers, the project "
     "demonstrates that a thoughtfully engineered, landmark-based, "
     "provider-abstracted architecture can make genuinely bilingual, bidirectional "
     "sign-language translation both practical and affordable.")

h2("15.2 Contributions Revisited")
bullet("A single web application delivering four translation flows for two languages "
       "with no installation and no special hardware.")
bullet("A landmark-based recognition pipeline pairing a TensorFlow-Lite ASL model "
       "with a custom INT8-quantised CNN–GRU Arabic model.")
bullet("Grammar-aware bidirectional gloss translation and a non-manual-marker "
       "annotation hook.")
bullet("A text-to-sign synthesis pipeline with indexed pgvector semantic search and "
       "a smoothed fingerspelling stitcher.")
bullet("A resilient provider abstraction with cloud-to-offline fallbacks.")
bullet("A production platform with JWT/Argon2id security, a real-time meeting mode, "
       "a bilingual RTL design system, and an end-to-end latency-optimisation phase.")

h2("15.3 Future Work")
h3("15.3.1 Continuous Sign-Language Recognition")
body("The most significant extension is moving from isolated-sign recognition "
     "debounced into sequences toward true continuous sign-language recognition that "
     "models co-articulation — the blending of one sign into the next — using "
     "sequence-to-sequence architectures such as Transformer encoders with "
     "connectionist temporal classification. This would remove the reliance on the "
     "debouncing buffer and handle natural, fluent signing.")
h3("15.3.2 Vocabulary Expansion")
body("Both vocabularies, and especially the 20-word Arabic set, should be expanded "
     "with more and more varied training data. The confusion analysis indicates that "
     "additional data, particularly for visually similar sign pairs, is the most "
     "direct route to higher accuracy.")
h3("15.3.3 Rendering Non-Manual Markers on the Avatar")
body("The non-manual-marker layer already classifies sentence type and emits a "
     "marker span; the natural next step is to make the avatar render those eyebrow, "
     "head, and mouth movements, completing the grammatical fidelity of the "
     "text-to-sign direction. The heuristic sentence-type classifier could also be "
     "replaced by a learned model behind the same interface.")
h3("15.3.4 On-Device and Edge Inference")
body("Running the recognition models entirely in the browser via WebAssembly or "
     "WebGPU would remove the landmark round-trip to the server, further reducing "
     "latency and strengthening the privacy guarantee, at the cost of a larger "
     "client download.")
h3("15.3.5 Personalisation and Dialectal Variation")
body("Sign languages vary regionally and individually. Allowing users to enrol a "
     "few examples of their own signing, or selecting a regional dialect, would "
     "improve accuracy for individual signers and broaden coverage of Egyptian and "
     "other Arabic sign-language variants.")
h3("15.3.6 Mobile and Native Clients")
body("While the web client is deliberately install-free, native mobile clients "
     "could exploit device cameras and on-device acceleration more fully and provide "
     "a more integrated experience, reusing the same server API and provider layer.")

h2("15.4 Closing Remarks")
body("Together began from a simple conviction: that the burden of bridging the "
     "communication gap should not fall entirely on Deaf and Hard-of-Hearing people, "
     "and that modern machine learning, carefully engineered, can carry part of that "
     "load. The system described in this thesis is a working, bilingual, "
     "bidirectional demonstration that such a bridge can be built with commodity "
     "hardware, in a browser, for free, and made resilient enough to keep working "
     "when the network does not. We hope it is a useful step, and that the "
     "architecture and findings documented here help others continue the work of "
     "bringing people together.")

# ═══════════════════════════ REFERENCES ═══════════════════════════
h1("References")
refs = [
    "World Health Organization. World Report on Hearing. Geneva: WHO, 2021.",
    "World Federation of the Deaf. Our Work and Sign Language Rights. https://wfdeaf.org.",
    "Stokoe, W. C. Sign Language Structure: An Outline of the Visual Communication Systems of the American Deaf. Studies in Linguistics, Occasional Papers 8, 1960.",
    "Sandler, W., and Lillo-Martin, D. Sign Language and Linguistic Universals. Cambridge University Press, 2006.",
    "Liddell, S. K. Grammar, Gesture, and Meaning in American Sign Language. Cambridge University Press, 2003.",
    "Lugaresi, C., et al. MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172, 2019.",
    "Google. MediaPipe Holistic: Simultaneous Face, Hand and Pose Prediction. Google AI Blog, 2020.",
    "LeCun, Y., Bengio, Y., and Hinton, G. Deep Learning. Nature, 521(7553):436–444, 2015.",
    "Hochreiter, S., and Schmidhuber, J. Long Short-Term Memory. Neural Computation, 9(8):1735–1780, 1997.",
    "Cho, K., et al. Learning Phrase Representations using RNN Encoder–Decoder for Statistical Machine Translation. EMNLP, 2014.",
    "Vaswani, A., et al. Attention Is All You Need. NeurIPS, 2017.",
    "Reimers, N., and Gurevych, I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. EMNLP-IJCNLP, 2019.",
    "Devlin, J., et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL-HLT, 2019.",
    "Li, D., et al. Word-level Deep Sign Language Recognition from Video: A New Large-scale Dataset and Methods Comparison (WLASL). WACV, 2020.",
    "Camgoz, N. C., et al. Neural Sign Language Translation. CVPR, 2018.",
    "Camgoz, N. C., et al. Sign Language Transformers: Joint End-to-End Sign Language Recognition and Translation. CVPR, 2020.",
    "Savitzky, A., and Golay, M. J. E. Smoothing and Differentiation of Data by Simplified Least Squares Procedures. Analytical Chemistry, 36(8):1627–1639, 1964.",
    "Maleh, Y., et al. Argon2: The Memory-Hard Function for Password Hashing and Other Applications (Password Hashing Competition winner). 2015.",
    "Jones, M., Bradley, J., and Sakimura, N. JSON Web Token (JWT). RFC 7519, IETF, 2015.",
    "Malkov, Y. A., and Yashunin, D. A. Efficient and Robust Approximate Nearest Neighbor Search Using Hierarchical Navigable Small World Graphs (HNSW). IEEE TPAMI, 2018.",
    "pgvector: Open-source vector similarity search for PostgreSQL. https://github.com/pgvector/pgvector.",
    "TensorFlow Lite / LiteRT and the XNNPACK delegate. Google, https://ai.google.dev/edge/litert.",
    "Paszke, A., et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library. NeurIPS, 2019.",
    "Radford, A., et al. Robust Speech Recognition via Large-Scale Weak Supervision (Whisper). OpenAI, 2022.",
    "FastAPI. https://fastapi.tiangolo.com. Socket.IO. https://socket.io. WebRTC. https://webrtc.org.",
    "Google. Gemini API documentation and pricing. https://ai.google.dev.",
    "Ollama: Run large language models locally. https://ollama.com.",
    "OWASP Foundation. Password Storage Cheat Sheet. https://cheatsheetseries.owasp.org.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"[{i}] ").bold = True
    p.add_run(r)

# ═══════════════════════════ APPENDICES ═══════════════════════════
h1("Appendix A: API Surface")
body("The following table summarises the public HTTP API and the real-time "
     "Socket.IO events of the Together server.")
table(
    ["Method & path", "Purpose"],
    [
        ["GET /, /login, /signup, /dashboard?lang=en|ar", "Server-rendered pages"],
        ["POST /api/auth/signup | login | refresh | logout", "Authentication"],
        ["GET /api/auth/me", "Current user profile"],
        ["POST /api/translate", "Landmarks → recognised gloss"],
        ["POST /api/translate/sentence", "Gloss sequence → natural sentence"],
        ["POST /api/gloss", "Sentence → gloss + non-manual markers"],
        ["GET /api/tts", "Text → speech audio"],
        ["POST /api/stt", "Speech audio → text"],
        ["GET /api/signs/lookup, /api/signs/{word}", "Single sign lookup"],
        ["POST /api/signs/batch, /api/signs_ar/batch", "Batched sign-landmark lookup"],
        ["GET /api/health, /api/metrics", "Health and latency profiling"],
        ["Socket.IO: join_room, leave_room, announce_presence", "Room presence"],
        ["Socket.IO: webrtc_offer/answer/ice_candidate", "WebRTC signalling"],
        ["Socket.IO: translate_sentence", "Live caption relay"],
    ],
    caption="The Together API surface.")

h1("Appendix B: Configuration Reference")
table(
    ["Variable", "Purpose", "Default"],
    [
        ["DATABASE_URL", "PostgreSQL + pgvector DSN", "postgresql+psycopg://…/together"],
        ["JWT_SECRET_KEY", "Token signing secret", "dev-insecure (override!)"],
        ["ACCESS_TOKEN_EXPIRE_MINUTES / REFRESH_TOKEN_EXPIRE_DAYS", "Token lifetimes", "60 / 7"],
        ["LLM_PROVIDER / TTS_PROVIDER / STT_PROVIDER", "Force a provider", "auto (cloud → offline)"],
        ["GEMINI_API_KEY, GEMINI_*_MODEL", "Gemini configuration", "—"],
        ["OLLAMA_HOST, OLLAMA_MODEL, WHISPER_MODEL", "Offline fallbacks", "—"],
        ["TFLITE_NUM_THREADS", "ASL interpreter threads", "min(4, cores)"],
        ["TORCH_INT8", "Enable Arabic INT8 quantization", "1 (on)"],
        ["ALLOWED_ORIGINS", "CORS allow-list", "localhost"],
        ["DB_POOL_SIZE, SLOW_REQUEST_MS", "Tuning", "—"],
    ],
    caption="Principal configuration variables.")

h1("Appendix C: Vocabularies")
body("The Arabic recognition model covers the following 20 signs:")
body("baby, eat, father, finish, good, happy, hear, house, important, love, mall, "
     "me, mosque, mother, normal, sad, stop, thanks, thinking, worry.")
body("The English ASL model covers a 250-sign vocabulary of common everyday words, "
     "stored as an index-to-sign label map alongside the model.")

h1("Appendix D: Selected Source-Code Listings")
body("This appendix reproduces representative, lightly trimmed excerpts of the "
     "Together source code referenced throughout the thesis, so that the "
     "implementation can be read alongside its description. The full source is "
     "maintained in the project repository.")

h2("D.1 Argon2id Hashing and Transparent Legacy Migration (auth.py)")
code_block([
    "_ph = PasswordHasher(time_cost=3, memory_cost=65536, parallelism=2,",
    "                     hash_len=32, salt_len=16)   # Argon2id (OWASP)",
    "",
    "def get_password_hash(password: str) -> str:",
    "    return _ph.hash(password)",
    "",
    "def verify_password(plain: str, hashed: str) -> bool:",
    "    # Verify against Argon2id or legacy bcrypt (transparent upgrade path).",
    "    if hashed.startswith(('$2b$', '$2a$', '$2y$')):",
    "        return bcrypt.checkpw(plain.encode(), hashed.encode())",
    "    try:",
    "        return _ph.verify(hashed, plain)",
    "    except VerifyMismatchError:",
    "        return False",
    "",
    "def needs_rehash(hashed: str) -> bool:",
    "    return hashed.startswith(('$2b$', '$2a$', '$2y$'))",
])

h2("D.2 Refresh-Token Rotation (auth.py)")
code_block([
    "def rotate_refresh_token(old_token_str: str, db) -> tuple[int, str]:",
    "    payload = jwt.decode(old_token_str, SECRET_KEY, algorithms=[ALGORITHM])",
    "    if payload.get('type') != 'refresh':",
    "        raise HTTPException(401, 'Not a refresh token.')",
    "    jti, user_id = payload.get('jti'), int(payload['sub'])",
    "    repo = RefreshTokenRepository(db)",
    "    rt = repo.get(jti)",
    "    if rt is None or rt.revoked:",
    "        raise HTTPException(401, 'Refresh token has been revoked.')",
    "    if rt.expires_at < datetime.utcnow():",
    "        raise HTTPException(401, 'Refresh token expired.')",
    "    repo.revoke(jti)                       # one-time use",
    "    _, new_signed = create_refresh_token(user_id, db)",
    "    return user_id, new_signed",
])

h2("D.3 Sliding-Window Rate Limiter (auth.py)")
code_block([
    "class _SlidingWindowLimiter:",
    "    def allow(self, key: str) -> bool:",
    "        now = time.time(); cutoff = now - self._window",
    "        with self._lock:",
    "            recent = [t for t in self._hits[key] if t > cutoff]",
    "            if len(recent) >= self._max:",
    "                self._hits[key] = recent; return False",
    "            recent.append(now); self._hits[key] = recent; return True",
    "",
    "_auth_limiter = _SlidingWindowLimiter(max_requests=10, window_seconds=60)",
])

h2("D.4 Indexed Semantic Sign Lookup (asl_service.py)")
code_block([
    "def match_word(self, query):",
    "    q = query.lower().strip(); key = q.replace(\"'\", '')",
    "    if q in self._word_set:   return q",
    "    if key in self._word_set: return key",
    "    if key in self.skip_semantic_match: return None",
    "    if key in self.exact_match_map:",
    "        fb = self.exact_match_map[key]",
    "        if fb in self._word_set: return fb",
    "    q_emb = self.model.encode(query)          # SBERT 384-d",
    "    hit = SignRepository(db).nearest(q_emb, self.LANGUAGE, self.MAX_DISTANCE)",
    "    return hit[0].word if hit else None       # indexed pgvector kNN",
])

h2("D.5 The CNN–GRU Arabic Model (sign_predictor.py)")
code_block([
    "class SignLanguageCNNGRU(nn.Module):",
    "    def __init__(self, input_dim=177, num_classes=20):",
    "        super().__init__()",
    "        self.conv1 = nn.Conv1d(input_dim, 128, 3, padding=1)",
    "        self.bn1   = nn.BatchNorm1d(128)",
    "        self.conv2 = nn.Conv1d(128, 128, 3, padding=1)",
    "        self.bn2   = nn.BatchNorm1d(128)",
    "        self.gru1  = nn.GRU(128, 64, num_layers=2, batch_first=True,",
    "                            bidirectional=True, dropout=0.3)",
    "        self.fc1, self.dropout = nn.Linear(128, 64), nn.Dropout(0.5)",
    "        self.fc2   = nn.Linear(64, num_classes)",
    "",
    "    def forward(self, x):",
    "        x = x.permute(0, 2, 1)",
    "        x = F.relu(self.bn1(self.conv1(x)))",
    "        x = F.relu(self.bn2(self.conv2(x)))",
    "        x = x.permute(0, 2, 1)",
    "        gru_out, _ = self.gru1(x)",
    "        out = F.relu(self.fc1(gru_out[:, -1, :]))",
    "        return self.fc2(self.dropout(out))",
])

h2("D.6 Stitching with Interpolated Bridges and Savitzky–Golay Smoothing (stitch.py)")
code_block([
    "def stitch_sequences(sequences, transition_frames=5,",
    "                     smooth_window=7, smooth_polyorder=2):",
    "    seqs = [np.asarray(s, np.float32) for s in sequences if s is not None]",
    "    if len(seqs) == 1: return seqs[0]",
    "    pieces, boundaries, cursor = [], [], 0",
    "    for i, seq in enumerate(seqs):",
    "        pieces.append(seq); cursor += seq.shape[0]",
    "        if i < len(seqs) - 1:",
    "            bridge = _interp_bridge(seq[-1], seqs[i+1][0], transition_frames)",
    "            pieces.append(bridge)",
    "            boundaries.append(cursor + bridge.shape[0] // 2)",
    "            cursor += bridge.shape[0]",
    "    combined = np.concatenate(pieces, axis=0)",
    "    return _smooth_boundaries(combined, boundaries, transition_frames,",
    "                              smooth_window, smooth_polyorder)",
])

h1("Appendix E: Deployment and Operations")
body("Together is packaged for reproducible deployment through Docker Compose, which "
     "brings up three services: PostgreSQL with the pgvector extension, the web "
     "application, and a local Ollama instance for the offline LLM fallback. The "
     "recommended first run copies the example environment file, supplies the "
     "necessary secrets, and builds the stack.")
code_block([
    "cp .env.example .env        # add GEMINI_API_KEY / JWT_SECRET_KEY, etc.",
    "docker compose up --build   # → http://localhost:8000",
    "# pull a local model for the offline LLM fallback:",
    "docker exec -it <ollama-container> ollama run llama3.2",
])
body("A manual deployment requires Python 3.11 or newer and a PostgreSQL 16 instance "
     "with the vector extension. The schema, including the pgvector extension and the "
     "HNSW index, is created by Alembic migrations, after which the server is started "
     "through a launcher that sets thread and machine-learning environment guards "
     "before importing the heavy libraries.")
code_block([
    "python -m venv .venv && source .venv/bin/activate",
    "pip install -r requirements.txt",
    "export DATABASE_URL='postgresql+psycopg://together:together@localhost:5432/together'",
    "export JWT_SECRET_KEY='change-me'",
    "alembic upgrade head        # pgvector extension + tables + HNSW index",
    "python start_server.py      # uvicorn on the Socket.IO ASGI app",
])
body("Operationally, the health endpoint supports an optional warm-up that loads and "
     "exercises the models so the first user request does not pay the cold-start "
     "cost, and the metrics endpoint exposes the per-stage latency profile described "
     "in Chapter 13 for live monitoring. Provider selection, token lifetimes, thread "
     "counts, and the CORS allow-list are all controlled through environment "
     "variables, so the same image serves development, cloud production, and "
     "air-gapped deployments without code changes.")

h1("Appendix F: Glossary of Terms")
gloss_terms = [
    ("Sign", "A single lexical unit of a sign language, defined by handshape, location, movement, orientation, and non-manual features."),
    ("Gloss", "An uppercase written label standing for a sign, used as an intermediate representation between signing and spoken language."),
    ("Topic–Comment", "A sentence structure, common in ASL, that establishes the topic first and then comments on it, rather than strict Subject–Verb–Object order."),
    ("Non-manual marker", "Grammatical information carried by the face, head, and body (e.g. eyebrow raises, head tilts) that spans manual signs."),
    ("Fingerspelling", "Spelling a word letter by letter using the manual alphabet, used for proper nouns and words without a dedicated sign."),
    ("Landmark", "A single tracked keypoint (an x, y, z coordinate) on the body, hands, or face, as produced by MediaPipe Holistic."),
    ("Holistic", "The MediaPipe pipeline that jointly estimates pose, hand, and face landmarks from one RGB frame."),
    ("Embedding", "A dense numeric vector representing the meaning of a word or phrase, such that similar meanings lie close together."),
    ("Cosine distance", "A measure of dissimilarity between two embedding vectors, equal to one minus their cosine similarity."),
    ("HNSW", "Hierarchical Navigable Small World, an index structure enabling fast approximate nearest-neighbour search over embeddings."),
    ("Quantization", "Representing model weights (and sometimes activations) with lower-precision integers to reduce size and speed up inference."),
    ("Debouncing", "Filtering a noisy stream of repeated predictions into a clean sequence of distinct events, here one gloss per held sign."),
    ("Provider chain", "An ordered list of interchangeable service adapters tried in turn, so a fallback runs when the preferred one is unavailable."),
    ("Signalling", "The exchange of connection metadata (offers, answers, ICE candidates) needed to establish a peer-to-peer WebRTC media session."),
]
table(["Term", "Definition"], gloss_terms)

# ───────────────────────── page numbers in footer ─────────────────────────
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "end")
run._r.append(fb); run._r.append(it); run._r.append(fs)

doc.save(OUT)
print("Saved", OUT)
print("Figures:", fig_count[0], "Tables:", tab_count[0])
